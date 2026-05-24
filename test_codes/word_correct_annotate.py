# word_correct_annotate.py

import torch
import re
from transformers import BertTokenizer, BertForMaskedLM
from docx import Document
from docx.shared import RGBColor
import os

# ⭐ 引入参考文献模块
from reference_checker import extract_references, check_reference

# ==================== 配置 ====================
MODEL_PATH = r"D:\py-90-day\Smart_text_checker\models\macbert4csc-base"
INPUT_WORD = r"D:\py-90-day\Smart_text_checker\test_files\test_error.docx"
OUTPUT_WORD = r"D:\py-90-day\Smart_text_checker\output\FINAL_SPLIT_VERSION.docx"

CONFIDENCE_THRESHOLD = 0.8

device = torch.device("cuda" if torch.cuda.is_available() else "cpu")

# ==================== 加载模型 ====================
print("🔍 加载模型...")
tokenizer = BertTokenizer.from_pretrained(MODEL_PATH)
model = BertForMaskedLM.from_pretrained(MODEL_PATH)
model.to(device)
model.eval()
print("✅ 模型加载完成")


# ==================== 句子切分 ====================
def split_sentences(text):
    parts = re.split(r'(。|！|？)', text)
    sentences = []
    for i in range(0, len(parts)-1, 2):
        sentences.append(parts[i] + parts[i+1])
    return sentences


# ==================== 纠错 ====================
def correct_sentence(sentence):
    inputs = tokenizer(sentence, return_tensors='pt', truncation=True, max_length=512).to(device)

    with torch.no_grad():
        outputs = model(**inputs)

    logits = outputs.logits[0]
    probs = torch.softmax(logits, dim=-1)

    pred_ids = torch.argmax(probs, dim=-1)
    confidence = probs.max(dim=-1).values

    pred_text = tokenizer.decode(pred_ids, skip_special_tokens=True).replace(' ', '')

    corrected = list(sentence)
    errors = []

    for i, ori_char in enumerate(sentence):
        if i >= len(pred_text):
            continue

        pred_char = pred_text[i]
        conf = confidence[i].item()

        if ori_char != pred_char and conf > CONFIDENCE_THRESHOLD and ori_char not in [' ', '\n', '\t']:
            corrected[i] = pred_char
            errors.append({
                "pos": i,
                "ori": ori_char,
                "new": pred_char,
                "conf": round(conf, 3)
            })

    return ''.join(corrected), errors


def correct_text(text):
    sentences = split_sentences(text)

    final_text = ""
    all_errors = []
    offset = 0

    for sent in sentences:
        corrected, errors = correct_sentence(sent)

        for e in errors:
            e["pos"] += offset
            all_errors.append(e)

        final_text += corrected
        offset += len(sent)

    return final_text, all_errors


# ==================== Word处理 ====================
doc = Document(INPUT_WORD)
paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

print("\n📝 开始正文纠错...")

# ===== 正文纠错 =====
for para in doc.paragraphs:
    text = para.text.strip()
    if not text:
        continue

    corrected, errors = correct_text(text)

    if errors:
        run1 = para.add_run(f"\n【原文】{text}")
        run1.font.color.rgb = RGBColor(255, 0, 0)

        run2 = para.add_run(f"\n【建议】{corrected}")
        run2.font.color.rgb = RGBColor(0, 128, 0)

        for e in errors:
            detail = f"\n → 位置{e['pos']}：{e['ori']}→{e['new']}（{e['conf']}）"
            run3 = para.add_run(detail)
            run3.font.color.rgb = RGBColor(0, 0, 255)

# ===== 参考文献检测 =====
print("\n📚 开始参考文献检测...")

refs = extract_references(paragraphs)

for ref in refs:
    errors = check_reference(ref)

    if errors:
        para = doc.add_paragraph()
        run = para.add_run(f"\n【参考文献错误】{ref}")
        run.font.color.rgb = RGBColor(255, 0, 0)

        for e in errors:
            run_e = para.add_run(f"\n → {e}")
            run_e.font.color.rgb = RGBColor(0, 0, 255)

# ===== 保存 =====
os.makedirs(os.path.dirname(OUTPUT_WORD), exist_ok=True)

try:
    doc.save(OUTPUT_WORD)
    print("\n🎉 完成：模块化系统运行成功")
except PermissionError:
    print("\n❌ 请关闭 Word 文件后再运行")