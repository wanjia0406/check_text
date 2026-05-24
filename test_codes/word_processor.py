from docx.shared import RGBColor

def annotate_paragraph(para, original, corrected, errors):
    run1 = para.add_run(f"\n【原文】{original}")
    run1.font.color.rgb = RGBColor(255, 0, 0)

    run2 = para.add_run(f"\n【建议】{corrected}")
    run2.font.color.rgb = RGBColor(0, 128, 0)

    for e in errors:
        run3 = para.add_run(
            f"\n → 位置{e['pos']}：{e['ori']}→{e['new']}（{e['conf']}）"
        )
        run3.font.color.rgb = RGBColor(0, 0, 255)