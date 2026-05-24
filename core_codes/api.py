import os
import zipfile
import torch
import uuid
import shutil
import re
import json

from fastapi import FastAPI, UploadFile, File, HTTPException, Form
from fastapi.responses import FileResponse, JSONResponse
from fastapi.middleware.cors import CORSMiddleware
from docx import Document
from docx.shared import RGBColor, Inches

from .nlp_corrector import TextCorrector
from .pdf_processor import extract_text_from_pdf, extract_pdf_structure
try:
    from .image_processor import extract_text_from_image, analyze_image_quality
    HAS_IMAGE_PROCESSOR = True
except ImportError:
    from .image_ocr import extract_text_from_image
    HAS_IMAGE_PROCESSOR = False

from .reference_checker import (
    extract_references,
    check_reference,
    suggest_reference_fix,
    check_index_sequence
)
from .image_table_checker import detect_image_table_errors, analyze_table_structure, analyze_table_content, apply_table_fixes_to_docx
from .reporter import generate_report, generate_detail_report

# === 文本相似度去重工具函数 ===

def calculate_similarity(str1, str2):
    """计算两个字符串的相似度（基于编辑距离）"""
    len1, len2 = len(str1), len(str2)
    if len1 == 0 and len2 == 0:
        return 1.0
    if len1 == 0 or len2 == 0:
        return 0.0
    
    # 初始化编辑距离矩阵
    dp = [[0] * (len2 + 1) for _ in range(len1 + 1)]
    for i in range(len1 + 1):
        dp[i][0] = i
    for j in range(len2 + 1):
        dp[0][j] = j
    
    for i in range(1, len1 + 1):
        for j in range(1, len2 + 1):
            if str1[i-1] == str2[j-1]:
                dp[i][j] = dp[i-1][j-1]
            else:
                dp[i][j] = min(dp[i-1][j], dp[i][j-1], dp[i-1][j-1]) + 1
    
    max_len = max(len1, len2)
    return 1.0 - (dp[len1][len2] / max_len)

def is_similar(text1, text2, threshold=0.8):
    """判断两个文本是否相似（相似度超过阈值）"""
    return calculate_similarity(text1.strip(), text2.strip()) >= threshold

def deduplicate_by_similarity(paragraphs, similarity_threshold=0.8):
    """基于相似度去重，保留第一次出现的版本"""
    unique_paragraphs = []
    for para in paragraphs:
        para_stripped = para.strip()
        if not para_stripped:
            continue
        
        # 检查是否与已保留的段落相似
        is_duplicate = False
        for unique_para in unique_paragraphs:
            if is_similar(para_stripped, unique_para.strip(), similarity_threshold):
                is_duplicate = True
                break
        
        if not is_duplicate:
            unique_paragraphs.append(para)
    
    return unique_paragraphs

# === END 文本相似度去重工具函数 ===

app = FastAPI(
    title="Smart Text Checker API",
    description="智能文本检测API - 支持错别字纠正、参考文献格式检测、图片OCR识别、图表检测",
    version="1.0.0"
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

MODEL_PATH = os.path.join(os.path.dirname(__file__), "..", "models", "macbert_finetuned")
device = torch.device("cuda" if torch.cuda.is_available() else "cpu")

# 全局变量，将在启动事件中初始化
corrector = None

# 添加启动事件，确保模型只加载一次
@app.on_event("startup")
async def startup_event():
    global corrector
    print("[INFO] 加载模型...")
    corrector = TextCorrector(MODEL_PATH, device)
    print("[OK] 模型加载完成")

@app.on_event("shutdown")
async def shutdown_event():
    global corrector
    print("[INFO] 清理模型资源...")
    corrector = None
    print("[OK] 资源清理完成")

SUPPORTED_EXTENSIONS = (".docx", ".pdf", ".png", ".jpg", ".jpeg", ".bmp", ".tiff", ".gif")
TEMP_DIR = "temp"

def get_docx_html(file_path):
    """提取Word文档的完整HTML内容，保留表格和图片"""
    try:
        import base64
        from xml.etree import ElementTree as ET

        html_parts = ['<div class="doc-content">']
        doc = Document(file_path)

        image_map = {}
        if file_path.endswith(".docx"):
            try:
                with zipfile.ZipFile(file_path, 'r') as z:
                    rels_path = "word/_rels/document.xml.rels"
                    if rels_path in z.namelist():
                        rels_content = z.read(rels_path).decode("utf-8")
                        for match in re.finditer(r'Id="(rId\d+)"[^>]*Target="media/([^"]+)"', rels_content):
                            rid = match.group(1)
                            img_name = match.group(2)
                            img_path = f"word/media/{img_name}"
                            if img_path in z.namelist():
                                img_data = z.read(img_path)
                                ext = img_name.split(".")[-1].lower()
                                mime_type = {
                                    "png": "image/png",
                                    "jpg": "image/jpeg",
                                    "jpeg": "image/jpeg",
                                    "gif": "image/gif",
                                    "bmp": "image/bmp"
                                }.get(ext, "image/png")
                                b64_data = base64.b64encode(img_data).decode("utf-8")
                                image_map[rid] = f"data:{mime_type};base64,{b64_data}"
            except Exception as img_err:
                print(f"[WARNING] 提取图片失败: {img_err}")

        for element in doc.element.body:
            tag_name = element.tag.split('}')[-1]

            if tag_name == 'p':
                para_content = []
                for child in element.iter():
                    if child.tag.endswith('}t') and child.text:
                        para_content.append(child.text)
                    elif child.tag.endswith('}drawing') or child.tag.endswith('}pict'):
                        for blip in child.iter():
                            if blip.tag.endswith('}blip'):
                                rid = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                                if rid and rid in image_map:
                                    img_src = image_map[rid]
                                    inline_style = ""
                                    for attrib in blip.attrib:
                                        if 'extent' in attrib.lower():
                                            pass
                                    img_tag = f'<br><img src="{img_src}" style="max-width:100%;height:auto;margin:8px 0;border:1px solid #ddd;border-radius:4px;">'
                                    para_content.append(img_tag)

                para_text = ''.join(para_content).strip()
                if para_text:
                    html_parts.append(f'<p>{para_text}</p>')

            elif tag_name == 'tbl':
                rows = []
                for row in element.iter():
                    if row.tag.endswith('}tr'):
                        cells = []
                        for cell in row.iter():
                            if cell.tag.endswith('}tc'):
                                cell_text = []
                                for t in cell.iter():
                                    if t.tag.endswith('}t') and t.text:
                                        cell_text.append(t.text)
                                cells.append(''.join(cell_text).strip())
                        if cells:
                            rows.append(f'<tr>{"".join(f"<td>{c}</td>" for c in cells)}</tr>')

                if rows:
                    html_parts.append(f'<table>{"".join(rows)}</table>')

        html_parts.append('</div>')
        return ''.join(html_parts)

    except Exception as e:
        print(f"[WARNING] 提取文档HTML失败: {e}")
        return ""


def _extract_pdf_body_text(pdf_path, pdf_data=None):
    """从PDF提取纯正文文本，排除表格行。先取 text_block，不足时用 text 补"""
    data = pdf_data if pdf_data is not None else extract_pdf_structure(pdf_path)
    tables_data = data.get("tables", [])
    table_row_set = set()
    table_cell_set = set()
    for tb in tables_data:
        for row in tb.get("content", []):
            row_text_norm = re.sub(r'\s+', ' ', " ".join(str(c).strip() for c in row if c and str(c).strip())).strip()
            if row_text_norm:
                table_row_set.add(row_text_norm)
            for c in row:
                cell_text = str(c).strip()
                if cell_text and len(cell_text) > 1:
                    table_cell_set.add(cell_text)

    def _is_table_fragment(text):
        text_norm = re.sub(r'\s+', ' ', text).strip()
        if text_norm in table_row_set:
            return True
        tokens = [t for t in text_norm.split() if len(t) > 1]
        if tokens:
            match_count = sum(1 for t in tokens if t in table_cell_set)
            if match_count >= len(tokens) * 0.5:
                return True
        text_compact = text_norm.replace(' ', '')
        if len(text_compact) >= 2:
            for cell in table_cell_set:
                cell_compact = cell.replace(' ', '')
                if len(cell_compact) >= 3 and (text_compact in cell_compact or cell_compact in text_compact):
                    return True
        return False

    # 先收集所有 text_block 的逐行内容
    block_content_set = set()
    for t in data.get("text", []):
        if t.get("type") == "text_block":
            content = t.get("content", "").strip()
            if content and not _is_table_fragment(content):
                for line in content.split("\n"):
                    stripped = line.strip()
                    if stripped and not _is_table_fragment(stripped):
                        block_content_set.add(stripped)

    paragraphs = []
    seen_paragraphs = set()
    for t in data.get("text", []):
        content = t.get("content", "").strip()
        if not content or content.startswith("【PDF表格") or content.startswith("【图片OCR"):
            continue
        ttype = t.get("type", "")
        if ttype == "text_block":
            # 逐行处理 text_block，过滤掉已出现在 text 中的行
            for line in content.split("\n"):
                stripped = line.strip()
                if not stripped:
                    continue
                if _is_table_fragment(stripped):
                    continue
                if stripped in seen_paragraphs:
                    continue
                seen_paragraphs.add(stripped)
                paragraphs.append(stripped)
        elif ttype == "text":
            # 处理 text 的逐行内容，过滤掉已出现在 text_block 中的行
            for line in content.split("\n"):
                stripped = line.strip()
                if not stripped:
                    continue
                if _is_table_fragment(stripped):
                    continue
                if stripped in seen_paragraphs:
                    continue
                # 只添加不在 text_block 内容中的行
                if stripped not in block_content_set:
                    seen_paragraphs.add(stripped)
                    paragraphs.append(stripped)
    return "\n".join(paragraphs), tables_data


def get_pdf_html(file_path, pdf_data=None):
    """提取PDF的完整HTML内容，保留表格结构和图片"""
    try:
        import base64
        html_parts = ['<div class="doc-content">']
        data = pdf_data if pdf_data is not None else extract_pdf_structure(file_path)

        images_by_page = {}
        for img in data.get("images", []):
            page = img.get("page", 1)
            if page not in images_by_page:
                images_by_page[page] = []
            images_by_page[page].append(img)

        tables_by_page = {}
        for tb in data.get("tables", []):
            page = tb.get("page", 1)
            if page not in tables_by_page:
                tables_by_page[page] = []
            tables_by_page[page].append(tb)

        texts_by_page = {}
        for t in data.get("text", []):
            page = t.get("page", 1)
            if page not in texts_by_page:
                texts_by_page[page] = []
            texts_by_page[page].append(t)

        tables_data = data.get("tables", [])
        table_row_set = set()
        table_cell_set = set()
        for tb in tables_data:
            for row in tb.get("content", []):
                row_text_norm = re.sub(r'\s+', ' ', " ".join(str(c).strip() for c in row if c and str(c).strip())).strip()
                if row_text_norm:
                    table_row_set.add(row_text_norm)
                for c in row:
                    cell_text = str(c).strip()
                    if cell_text and len(cell_text) > 1:
                        table_cell_set.add(cell_text)

        def _is_table_fragment(text):
            text_norm = re.sub(r'\s+', ' ', text).strip()
            if text_norm in table_row_set:
                return True
            tokens = [t for t in text_norm.split() if len(t) > 1]
            if tokens:
                match_count = sum(1 for t in tokens if t in table_cell_set)
                if match_count >= len(tokens) * 0.5:
                    return True
            text_compact = text_norm.replace(' ', '')
            if len(text_compact) >= 2:
                for cell in table_cell_set:
                    cell_compact = cell.replace(' ', '')
                    if len(cell_compact) >= 3 and (text_compact in cell_compact or cell_compact in text_compact):
                        return True
            return False

        # 先收集所有 text_block 的逐行内容（用于精确去重）
        block_content_set = set()
        for t in data.get("text", []):
            if t.get("type") == "text_block":
                content = t.get("content", "").strip()
                if content and not _is_table_fragment(content):
                    for line in content.split("\n"):
                        stripped = line.strip()
                        if stripped and not _is_table_fragment(stripped):
                            block_content_set.add(stripped)

        all_pages = sorted(set(
            list(images_by_page.keys()) + list(tables_by_page.keys()) + list(texts_by_page.keys())
        ))

        seen_lines = set()

        for page in all_pages:
            elements = []
            page_texts = texts_by_page.get(page, [])
            page_tables = tables_by_page.get(page, [])
            page_images = images_by_page.get(page, [])

            for img in page_images:
                y = (img.get("position") or {}).get("y0", 0)
                elements.append((y, "image", img))

            for tb in page_tables:
                y = (tb.get("bbox") or {}).get("y0", 0)
                elements.append((y, "table", tb))

            for t in page_texts:
                content = t.get("content", "").strip()
                if not content:
                    continue
                if content.startswith("【PDF表格") or content.startswith("【图片OCR"):
                    continue

                ttype = t.get("type", "")
                if ttype == "text_block":
                    if _is_table_fragment(content):
                        continue
                    y = (t.get("position") or {}).get("y0", 0)
                    elements.append((y, "text", {"content": content, "type": "text_block"}))
                elif ttype == "text":
                    # 逐行处理 text，过滤掉已出现在 text_block 中的行
                    for idx, line in enumerate(content.split("\n")):
                        stripped = line.strip()
                        if not stripped:
                            continue
                        if _is_table_fragment(stripped):
                            continue
                        if stripped in seen_lines:
                            continue
                        # 只添加不在 text_block 内容中的行
                        if stripped not in block_content_set:
                            seen_lines.add(stripped)
                            y = idx * 10
                            text_data = {"content": stripped, "type": "full_text_line"}
                            elements.append((y, "text", text_data))

            # 按位置排序元素
            elements.sort(key=lambda x: x[0])

            # 按顺序添加到HTML
            for _, etype, edata in elements:
                if etype == "image":
                    img_bytes = edata.get("image")
                    ext = edata.get("extension", "png")
                    if img_bytes:
                        mime = {"png": "image/png", "jpg": "image/jpeg", "jpeg": "image/jpeg",
                                "gif": "image/gif", "bmp": "image/bmp"}.get(ext, "image/png")
                        b64 = base64.b64encode(img_bytes).decode("utf-8")
                        html_parts.append(
                            f'<img src="data:{mime};base64,{b64}" '
                            f'style="max-width:100%;height:auto;margin:8px 0;border:1px solid #ddd;border-radius:4px;">'
                        )

                elif etype == "table":
                    rows = edata.get("content", [])
                    if not rows:
                        continue
                    html_parts.append(
                        '<table style="width:100%;border-collapse:collapse;margin:1rem 0;'
                        'font-size:0.9rem;border:1px solid #666;">'
                    )
                    for i, row in enumerate(rows):
                        tag = "th" if i == 0 else "td"
                        bg_style = "background:#f5f5f5;font-weight:600;" if i == 0 else ""
                        cells_html = "".join(
                            f'<{tag} style="border:1px solid #999;padding:6px 8px;{bg_style}">{cell}</{tag}>'
                            for cell in row
                        )
                        html_parts.append(f"<tr>{cells_html}</tr>")
                    html_parts.append("</table>")

                elif etype == "text":
                    content = edata.get("content", "").strip()
                    if not content:
                        continue
                    if content.startswith("【PDF表格") or content.startswith("【图片OCR"):
                        continue
                    html_parts.append(f"<p>{content}</p>")

        html_parts.append("</div>")
        return "".join(html_parts)
    except Exception as e:
        print(f"[WARNING] 生成PDF HTML失败: {e}")
        import traceback
        traceback.print_exc()
        return ""


def build_pdf_corrected_docx(pdf_path, correct_text=True):
    """构建PDF修正文档，保持文本/图片/表格的原始页面位置顺序"""
    try:
        data = extract_pdf_structure(pdf_path)
        doc = Document()

        images_by_page = {}
        for img in data.get("images", []):
            page = img.get("page", 1)
            images_by_page.setdefault(page, []).append(img)

        tables_by_page = {}
        for tb in data.get("tables", []):
            page = tb.get("page", 1)
            tables_by_page.setdefault(page, []).append(tb)

        texts_by_page = {}
        for t in data.get("text", []):
            page = t.get("page", 1)
            texts_by_page.setdefault(page, []).append(t)

        tables_data = data.get("tables", [])
        table_row_set = set()
        table_cell_set = set()
        for tb in tables_data:
            for row in tb.get("content", []):
                row_text_norm = re.sub(r'\s+', ' ', " ".join(str(c).strip() for c in row if c and str(c).strip())).strip()
                if row_text_norm:
                    table_row_set.add(row_text_norm)
                for c in row:
                    cell_text = str(c).strip()
                    if cell_text and len(cell_text) > 1:
                        table_cell_set.add(cell_text)

        def _is_table_fragment(text):
            text_norm = re.sub(r'\s+', ' ', text).strip()
            if text_norm in table_row_set:
                return True
            tokens = [t for t in text_norm.split() if len(t) > 1]
            if tokens:
                match_count = sum(1 for t in tokens if t in table_cell_set)
                if match_count >= len(tokens) * 0.5:
                    return True
            text_compact = text_norm.replace(' ', '')
            if len(text_compact) >= 2:
                for cell in table_cell_set:
                    cell_compact = cell.replace(' ', '')
                    if len(cell_compact) >= 3 and (text_compact in cell_compact or cell_compact in text_compact):
                        return True
            return False

        # 先收集所有 text_block 的逐行内容（用于精确去重）
        block_content_set = set()
        for t in data.get("text", []):
            if t.get("type") == "text_block":
                content = t.get("content", "").strip()
                if content and not _is_table_fragment(content):
                    for line in content.split("\n"):
                        stripped = line.strip()
                        if stripped and not _is_table_fragment(stripped):
                            block_content_set.add(stripped)

        all_pages = sorted(set(
            list(images_by_page.keys()) + list(tables_by_page.keys()) + list(texts_by_page.keys())
        ))

        seen_lines = set()

        for page in all_pages:
            elements = []
            page_texts = texts_by_page.get(page, [])
            page_tables = tables_by_page.get(page, [])
            page_images = images_by_page.get(page, [])

            for img in page_images:
                y = (img.get("position") or {}).get("y0", 0)
                elements.append((y, "image", img))

            for tb in page_tables:
                y = (tb.get("bbox") or {}).get("y0", 0)
                elements.append((y, "table", tb))

            for t in page_texts:
                content = t.get("content", "").strip()
                if not content:
                    continue
                if content.startswith("【PDF表格") or content.startswith("【图片OCR"):
                    continue

                ttype = t.get("type", "")
                if ttype == "text_block":
                    if _is_table_fragment(content):
                        continue
                    y = (t.get("position") or {}).get("y0", 0)
                    elements.append((y, "text", {"content": content, "type": "text_block"}))
                elif ttype == "text":
                    # 逐行处理 text，过滤掉已出现在 text_block 中的行
                    for idx, line in enumerate(content.split("\n")):
                        stripped = line.strip()
                        if not stripped:
                            continue
                        if _is_table_fragment(stripped):
                            continue
                        if stripped in seen_lines:
                            continue
                        # 只添加不在 text_block 内容中的行
                        if stripped not in block_content_set:
                            seen_lines.add(stripped)
                            y = idx * 10
                            text_data = {"content": stripped, "type": "full_text_line"}
                            elements.append((y, "text", text_data))

            # 按位置排序元素
            elements.sort(key=lambda x: x[0])

            # 按顺序添加到文档
            prev_type = None
            for _, etype, edata in elements:
                if etype == "table" and prev_type == "table":
                    doc.add_paragraph('')
                prev_type = etype

                if etype == "image":
                    img_bytes = edata.get("image")
                    ext = edata.get("extension", "png")
                    if img_bytes:
                        temp_img_path = os.path.join(TEMP_DIR, f"_pdf_img_{uuid.uuid4().hex}.{ext}")
                        try:
                            with open(temp_img_path, "wb") as f_img:
                                f_img.write(img_bytes)
                            doc.add_picture(temp_img_path, width=Inches(5.5))
                            last_paragraph = doc.paragraphs[-1]
                            last_paragraph.alignment = 1
                        finally:
                            try:
                                os.unlink(temp_img_path)
                            except:
                                pass

                elif etype == "table":
                    rows = edata.get("content", [])
                    if not rows:
                        continue
                    table = doc.add_table(rows=len(rows), cols=len(rows[0]) if rows else 1)
                    table.style = 'Table Grid'
                    for i, row_data in enumerate(rows):
                        for j, cell_text in enumerate(row_data):
                            if j < len(table.rows[i].cells):
                                table.rows[i].cells[j].text = cell_text

                elif etype == "text":
                    content = edata.get("content", "").strip()
                    if not content:
                        continue
                    if content.startswith("【PDF表格") or content.startswith("【图片OCR"):
                        continue
                    if correct_text:
                        try:
                            corrected_content, _ = corrector.correct_text(content)
                            # 对参考文献内容进行修正
                            refs = extract_references([content])
                            for ref in refs:
                                errors = check_reference(ref)
                                if errors:
                                    fixed_ref = suggest_reference_fix(ref)
                                    if fixed_ref != ref:
                                        corrected_content = corrected_content.replace(ref, fixed_ref)
                            doc.add_paragraph(corrected_content if corrected_content else content)
                        except Exception:
                            doc.add_paragraph(content)
                    else:
                        doc.add_paragraph(content)

        # 对整个文档的段落进行参考文献修正
        for para in doc.paragraphs:
            if para.text.strip():
                refs = extract_references([para.text])
                for ref in refs:
                    errors = check_reference(ref)
                    if errors:
                        fixed_ref = suggest_reference_fix(ref)
                        if fixed_ref != ref:
                            para.text = para.text.replace(ref, fixed_ref)

        apply_table_fixes_to_docx(doc, corrector)

        print(f"[DEBUG build_pdf_docx] FINAL: {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables")
        return doc
    except Exception as e:
        print(f"[WARNING] 构建PDF修正文档失败: {e}")
        import traceback as _tb
        _tb.print_exc()
        return None


def read_file(file_path):
    """
    读取文件内容，支持docx、pdf和图片格式
    
    Args:
        file_path: 文件路径
        
    Returns:
        tuple: (doc对象, 段落列表, 文件类型, 表格内容列表, [图片质量信息])
    """
    if file_path.endswith(".docx"):
        doc = Document(file_path)
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        
        # 提取表格内容
        tables_content = []
        for table in doc.tables:
            table_data = {
                'rows': [],
                'header': [],
                'row_count': len(table.rows),
                'col_count': len(table.columns) if table.rows else 0
            }
            for i, row in enumerate(table.rows):
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text.strip())
                table_data['rows'].append(row_data)
                if i == 0:
                    table_data['header'] = row_data
            tables_content.append(table_data)
        
        return doc, paragraphs, "docx", tables_content

    elif file_path.endswith(".pdf"):
        text, tables_data = _extract_pdf_body_text(file_path)
        paragraphs = [t.strip() for t in text.split("\n") if t.strip()]
        doc = Document()

        tables_content = []
        for tb in tables_data:
            content = tb.get("content", [])
            if content:
                table_data = {
                    "rows": content,
                    "header": content[0] if content else [],
                    "row_count": len(content),
                    "col_count": len(content[0]) if content else 0,
                    "page": tb.get("page", 1)
                }
                tables_content.append(table_data)

        return doc, paragraphs, "pdf", tables_content

    elif file_path.endswith((".png", ".jpg", ".bmp", ".tiff", ".gif")):
        text = extract_text_from_image(file_path)
        # 使用相似度去重（阈值80%）
        paragraphs = text.split("\n")
        paragraphs = deduplicate_by_similarity(paragraphs, similarity_threshold=0.8)
        doc = Document()
        quality_info = {}
        if HAS_IMAGE_PROCESSOR:
            quality_info = analyze_image_quality(file_path)
        return doc, paragraphs, "image", [], quality_info

    else:
        raise ValueError(f"不支持的文件类型，支持的类型: {SUPPORTED_EXTENSIONS}")

def extract_images_from_docx(docx_path, output_dir="temp_images"):
    """
    从Word文档中提取图片
    
    Args:
        docx_path: Word文档路径
        output_dir: 图片输出目录
        
    Returns:
        list: 提取的图片路径列表
    """
    os.makedirs(output_dir, exist_ok=True)
    image_paths = []

    try:
        with zipfile.ZipFile(docx_path, 'r') as zip_ref:
            for file in zip_ref.namelist():
                if file.startswith("word/media/"):
                    zip_ref.extract(file, output_dir)

        media_path = os.path.join(output_dir, "word", "media")
        if os.path.exists(media_path):
            for img in os.listdir(media_path):
                image_paths.append(os.path.join(media_path, img))

    except Exception as e:
        print(f"[WARNING] 提取图片失败: {e}")

    return image_paths

def generate_corrected_doc(file_path, output_path):
    """生成修正后的文档，保留原始文档的图片和表格"""
    try:
        import shutil
        shutil.copy(file_path, output_path)
        
        doc = Document(output_path)
        
        for para in doc.paragraphs:
            if para.text.strip():
                corrected, errors = corrector.correct_text(para.text)
                if corrected and corrected != para.text:
                    para.text = corrected
        
        apply_table_fixes_to_docx(doc, corrector)

        paragraphs = [p.text for p in doc.paragraphs]
        refs = extract_references(paragraphs)
        
        for ref in refs:
            errors = check_reference(ref)
            if errors:
                fixed_ref = suggest_reference_fix(ref)
                if fixed_ref != ref:
                    for para in doc.paragraphs:
                        if ref in para.text:
                            para.text = para.text.replace(ref, fixed_ref)
        
        doc.save(output_path)
        return True
        
    except Exception as e:
        print(f"[WARNING] 生成修正文档失败: {e}")
        try:
            result = read_file(file_path)
            if len(result) == 4:
                doc, paragraphs, file_type, tables_content = result
            else:
                doc, paragraphs, file_type = result
        except ValueError as e:
            raise HTTPException(status_code=400, detail=str(e))

        doc_new = Document()

        for text in paragraphs:
            if not text:
                continue

            corrected, errors = corrector.correct_text(text)
            doc_new.add_paragraph(corrected if corrected else text)

        # 将表格内容也写入修正文档
        if len(result) >= 4 and result[3]:
            for tbl in result[3]:
                rows = tbl.get("rows", [])
                if not rows:
                    continue
                table = doc_new.add_table(rows=len(rows), cols=len(rows[0]) if rows else 1)
                table.style = 'Table Grid'
                for i, row_data in enumerate(rows):
                    for j, cell_text in enumerate(row_data):
                        if j < len(table.rows[i].cells):
                            # 对表格内容纠错
                            corrected_cell, _ = corrector.correct_text(cell_text)
                            table.rows[i].cells[j].text = corrected_cell if corrected_cell else cell_text

        doc_new.save(output_path)
        return True

def copy_to_corrected_doc(input_path, output_path):
    """直接复制检测结果文档作为修正文档（避免重复处理）"""
    import shutil
    # 由于process_file已经进行过文本纠错，这里直接复制output文档
    # output文档包含纠错后的纯文本内容
    output_doc_path = output_path.replace("corrected_", "output_")
    if os.path.exists(output_doc_path):
        shutil.copy(output_doc_path, output_path)
        return True
    return generate_corrected_doc(input_path, output_path)

def process_file(file_path, output_path, add_annotations=True, corrected_output_path=None):
    """
    处理文档，执行文本纠错、参考文献检测、图片表格检测等
    
    Args:
        file_path: 输入文件路径
        output_path: 检测结果文档输出路径
        add_annotations: 是否添加批注（默认True）
        corrected_output_path: 纯修正文档输出路径
        
    Returns:
        list: 错误列表
    """
    # 检查模型是否已初始化
    if corrector is None:
        raise RuntimeError("模型尚未初始化，请确保服务已正确启动")
    
    try:
        result = read_file(file_path)
        if len(result) == 5:
            doc, paragraphs, file_type, tables_content, image_quality = result
        elif len(result) == 4:
            doc, paragraphs, file_type, tables_content = result
            image_quality = {}
        else:
            doc, paragraphs, file_type = result
            tables_content = []
            image_quality = {}
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))

    # 如果需要同时生成纯修正文档，先复制原始文件
    if corrected_output_path and file_type == "docx":
        shutil.copy(file_path, corrected_output_path)
        doc_corrected = Document(corrected_output_path)
        
        for para in doc_corrected.paragraphs:
            if para.text.strip():
                corrected, errors = corrector.correct_text(para.text)
                if corrected and corrected != para.text:
                    para.text = corrected
        
        apply_table_fixes_to_docx(doc_corrected, corrector)
    else:
        doc_corrected = None
        if corrected_output_path:
            doc_corrected = Document()

    all_errors = []
    all_text = "\n".join(paragraphs)

    table_cell_set = set()
    if file_type != "docx" and tables_content:
        for tbl in tables_content:
            for row in tbl.get("rows", []):
                for cell in row:
                    cell_text = str(cell).strip()
                    if cell_text and len(cell_text) > 1:
                        table_cell_set.add(cell_text)

    print("[INFO] 正在进行文本纠错...")
    for text in paragraphs:
        if not text:
            continue

        corrected, errors = corrector.correct_text(text)

        if add_annotations:
            if errors:
                para = doc.add_paragraph()
                para.add_run(f"【原文】{text}").font.color.rgb = RGBColor(255, 0, 0)
                para.add_run(f"\n【建议】{corrected}").font.color.rgb = RGBColor(0, 128, 0)

                for e in errors:
                    para.add_run(f"\n → {e['text']} → {e['suggestion']}").font.color.rgb = RGBColor(0, 0, 255)
                all_errors.extend(errors)
            else:
                doc.add_paragraph(text)
        else:
            doc.add_paragraph(corrected if corrected else text)

        # 同时生成纯修正文档（复用纠错结果）- 仅用于非docx格式
        if doc_corrected and file_type != "docx":
            stripped = text.strip()
            if stripped and table_cell_set:
                fragments = [s.strip() for s in stripped.replace("\t", " ").split() if len(s.strip()) > 1]
                if stripped in table_cell_set:
                    continue
                if fragments and len(fragments) >= 3:
                    match_count = sum(1 for f in fragments if f in table_cell_set)
                    if match_count >= len(fragments) * 0.5:
                        continue
            doc_corrected.add_paragraph(corrected if corrected else text)

    # 将表格内容写入纯修正文档（PDF等非docx格式）
    if doc_corrected and file_type != "docx" and tables_content:
        for tbl in tables_content:
            rows = tbl.get("rows", [])
            if not rows:
                continue
            table = doc_corrected.add_table(rows=len(rows), cols=len(rows[0]) if rows else 1)
            table.style = 'Table Grid'
            for i, row_data in enumerate(rows):
                for j, cell_text in enumerate(row_data):
                    if j < len(table.rows[i].cells):
                        corrected_cell, _ = corrector.correct_text(cell_text)
                        table.rows[i].cells[j].text = corrected_cell if corrected_cell else cell_text

    # 处理表格内容的文本纠错（使用 image_table_checker 中的专用函数）
    if tables_content:
        print(f"[INFO] 正在处理 {len(tables_content)} 个表格的文本纠错...")
        for table_idx, table_data in enumerate(tables_content, 1):
            table_data_with_index = table_data.copy()
            table_data_with_index['index'] = table_idx
            table_errors = analyze_table_content(table_data_with_index, corrector)
            all_errors.extend(table_errors)

    if file_type == "docx":
        print("[INFO] 正在处理Word文档中的图片...")
        image_paths = extract_images_from_docx(file_path)
        print(f"[INFO] 发现 {len(image_paths)} 张图片，跳过图片OCR处理")

    if file_type == "image":
        print("[INFO] 正在分析图片质量...")
        if image_quality:
            para = doc.add_paragraph()
            para.add_run("【图片质量分析】\n").bold = True
            para.add_run(f"质量等级: {image_quality.get('quality_level', '未知')}\n")
            
            if "issues" in image_quality and image_quality["issues"]:
                para.add_run(f"问题: {', '.join(image_quality['issues'])}\n")
            
            if "suggestion" in image_quality:
                para.add_run(f"建议: {image_quality['suggestion']}\n")

    print("[INFO] 正在检测图片表格相关问题...")
    table_errors = detect_image_table_errors(all_text)
    all_errors.extend(table_errors)

    # 注意：图片表格错误会在最后的统一报告中添加，不需要在这里重复添加

    print("[INFO] 正在检测参考文献...")
    refs = extract_references(paragraphs)
    
    # 去重处理，避免重复计数
    seen_refs = set()
    unique_refs = []
    for ref in refs:
        if ref not in seen_refs:
            seen_refs.add(ref)
            unique_refs.append(ref)
    print(f"[INFO] 原始参考文献数: {len(refs)}, 去重后: {len(unique_refs)}")
    
    refs = unique_refs

    index_errors = check_index_sequence(refs)
    for err in index_errors:
        all_errors.append({
            "type": "reference",
            "level": "error",
            "text": "参考文献编号",
            "message": err
        })

    for ref in refs:
        errors = check_reference(ref)
        all_errors.extend(errors)

        if add_annotations:
            if errors:
                para = doc.add_paragraph()
                para.add_run(f"【参考文献错误】{ref}").font.color.rgb = RGBColor(255, 0, 0)

                for e in errors:
                    para.add_run(f"\n → {e['message']}").font.color.rgb = RGBColor(0, 0, 255)

                para.add_run(f"\n【建议】{suggest_reference_fix(ref)}").font.color.rgb = RGBColor(0, 128, 0)
        else:
            if errors:
                fixed_ref = suggest_reference_fix(ref)
                doc.add_paragraph(fixed_ref)
            else:
                doc.add_paragraph(ref)
        
        # 在修正文档中也处理参考文献
        if doc_corrected and file_type == "docx":
            fixed_ref = suggest_reference_fix(ref) if errors else ref
            for para in doc_corrected.paragraphs:
                if ref in para.text and ref != fixed_ref:
                    para.text = para.text.replace(ref, fixed_ref)

    print("[INFO] 正在生成统一错误报告...")
    if all_errors and add_annotations:
        # 去重处理，避免报告中出现重复错误
        seen_errors = set()
        unique_errors = []
        for err in all_errors:
            # 使用错误文本和类型作为唯一标识
            error_key = (err.get('text', ''), err.get('type', 'other'))
            if error_key not in seen_errors:
                seen_errors.add(error_key)
                unique_errors.append(err)

        doc.add_page_break()

        report_para = doc.add_paragraph()
        report_para.add_run("【智能检测报告 - 错误汇总】\n").bold = True
        report_para.add_run("=" * 50 + "\n")

        error_groups = {}
        for err in unique_errors:
            err_type = err.get('type', 'other')
            if err_type not in error_groups:
                error_groups[err_type] = []
            error_groups[err_type].append(err)

        for err_type, errors in error_groups.items():
            type_names = {
                'spelling': '错别字',
                'grammar': '语法问题',
                'semantic': '语义问题',
                'reference': '参考文献格式',
                'image': '图片问题',
                'table': '表格问题',
                'other': '其他问题'
            }
            
            report_para.add_run(f"\n【{type_names.get(err_type, err_type)}】共 {len(errors)} 个问题\n")
            report_para.add_run("-" * 30 + "\n")

            for i, err in enumerate(errors, 1):
                original = err.get('text', '')
                suggestion = err.get('suggestion', '')
                message = err.get('message', '')
                
                if original and suggestion:
                    report_para.add_run(f"  {i}. 原文: {original} → {suggestion}\n")
                elif message:
                    report_para.add_run(f"  {i}. {message}\n")

    output_dir = os.path.dirname(output_path)
    if output_dir and not os.path.exists(output_dir):
        os.makedirs(output_dir, exist_ok=True)
    
    doc.save(output_path)
    print("[INFO] 文件处理完成")

    # 保存纯修正文档
    if doc_corrected and corrected_output_path:
        if file_type == "pdf":
            doc_corrected = build_pdf_corrected_docx(file_path)

        if doc_corrected:
            doc_corrected.save(corrected_output_path)
            print("[INFO] 纯修正文档生成完成")

    return all_errors

@app.post("/upload", response_description="文件上传并检测")
async def upload(file: UploadFile = File(...)):
    file_ext = os.path.splitext(file.filename)[1].lower()
    if file_ext not in SUPPORTED_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail=f"不支持的文件类型: {file_ext}。支持的类型: {SUPPORTED_EXTENSIONS}"
        )

    file_id = str(uuid.uuid4())
    os.makedirs(TEMP_DIR, exist_ok=True)

    input_path = os.path.join(TEMP_DIR, f"temp_{file_id}_{file.filename}")
    output_path = os.path.join(TEMP_DIR, f"output_{file_id}.docx")

    try:
        with open(input_path, "wb") as f:
            shutil.copyfileobj(file.file, f)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"文件保存失败: {str(e)}")

    # 同时生成检测结果文档和纯修正文档
    corrected_output_path = os.path.join(TEMP_DIR, f"corrected_{file_id}_{file.filename}")
    try:
        errors = process_file(input_path, output_path, corrected_output_path=corrected_output_path)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"文件处理失败: {str(e)}")

    report = generate_report(errors)
    details = generate_detail_report(errors)

    error_stats = {}
    for err in errors:
        err_type = err.get('type', 'other')
        error_stats[err_type] = error_stats.get(err_type, 0) + 1

    # 获取原始文本内容（段落 + 表格内容，确保前端预览能高亮表格内部错误）
    original_text = ""
    original_html = ""
    tables = []
    images = []
    try:
        result = read_file(input_path)
        if len(result) >= 2:
            paragraphs = result[1]
            parts = list(paragraphs)
            # 表格内容也加入文本，确保前端文本预览路径也能匹配表格内部错误
            if len(result) >= 4:
                tables_content = result[3]
                for idx, tbl in enumerate(tables_content, 1):
                    for row in tbl.get('rows', []):
                        for cell in row:
                            if cell and cell.strip():
                                parts.append(cell.strip())
                    tables.append({
                        "index": idx,
                        "rows": tbl.get("rows", []),
                        "row_count": tbl.get("row_count", 0),
                        "col_count": tbl.get("col_count", 0),
                        "header": tbl.get("header", []),
                        "caption": f"表格{idx}"
                    })
            original_text = "\n".join(parts)
        if input_path.endswith(".docx"):
            original_html = get_docx_html(input_path)
        elif input_path.endswith(".pdf"):
            import base64 as _b64
            pdf_data = extract_pdf_structure(input_path)
            original_html = get_pdf_html(input_path, pdf_data=pdf_data)
            # 从PDF提取图片，转为base64返回给前端
            try:
                for idx, img in enumerate(pdf_data.get("images", []), 1):
                    img_bytes = img.get("image")
                    ext = img.get("extension", "png")
                    if img_bytes:
                        mime = {"png": "image/png", "jpg": "image/jpeg", "jpeg": "image/jpeg",
                                "gif": "image/gif", "bmp": "image/bmp"}.get(ext, "image/png")
                        images.append({
                            "index": idx,
                            "src": f"data:{mime};base64,{_b64.b64encode(img_bytes).decode('utf-8')}",
                            "caption": f"PDF图片{idx}（第{img.get('page', '?')}页）",
                            "position": f"第{img.get('page', '?')}页"
                        })
            except Exception as e:
                print(f"[WARNING] 提取PDF图片失败: {e}")
        elif input_path.endswith((".png", ".jpg", ".jpeg", ".bmp", ".tiff", ".gif")):
            # 处理纯图片文件，显示提取的文本内容（类似Word/PDF预览）
            import base64 as _b64
            try:
                with open(input_path, "rb") as f:
                    img_bytes = f.read()
                ext = input_path.split(".")[-1].lower()
                mime = {"png": "image/png", "jpg": "image/jpeg", "jpeg": "image/jpeg",
                        "gif": "image/gif", "bmp": "image/bmp", "tiff": "image/tiff"}.get(ext, "image/png")
                
                images.append({
                    "index": 1,
                    "src": f"data:{mime};base64,{_b64.b64encode(img_bytes).decode('utf-8')}",
                    "caption": f"图片文件",
                    "position": "第1页"
                })
                
                # 生成HTML展示提取的文本内容（类似Word/PDF预览）
                if paragraphs:
                    html_parts = ['<div class="doc-content">']
                    for idx, para in enumerate(paragraphs, 1):
                        if para.strip():
                            html_parts.append(f'<p class="paragraph" data-index="{idx}">{para}</p>')
                    html_parts.append('</div>')
                    original_html = '\n'.join(html_parts)
                else:
                    # 如果没有提取到文本，显示图片
                    original_html = f'<div class="doc-content"><img src="data:{mime};base64,{_b64.b64encode(img_bytes).decode("utf-8")}" style="max-width:100%;height:auto;"></div>'
            except Exception as e:
                print(f"[WARNING] 处理图片文件失败: {e}")
    except Exception as e:
        print(f"[WARNING] 获取原始文本失败: {e}")

    return {
        "file_id": file_id,
        "original_filename": file.filename,
        "report": report,
        "details": details,
        "errors": errors,
        "error_stats": error_stats,
        "total_errors": len(errors),
        "download_url": f"/download/{os.path.basename(output_path)}",
        "original_text": original_text,
        "original_html": original_html,
        "tables": tables,
        "images": images
    }

@app.get("/download/{filename}", response_description="下载检测结果文件")
def download(filename: str):
    if '..' in filename or '/' in filename or '\\' in filename:
        raise HTTPException(status_code=400, detail="无效的文件名")
    
    file_path = os.path.join(TEMP_DIR, filename)

    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="文件不存在")

    return FileResponse(
        path=file_path,
        filename="result.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

@app.get("/download/corrected/{file_id}/{original_filename}", response_description="下载纯修复文档")
def download_corrected(file_id: str, original_filename: str):
    if '..' in file_id or '/' in file_id or '\\' in file_id:
        raise HTTPException(status_code=400, detail="无效的文件ID")
    
    output_path = os.path.join(TEMP_DIR, f"corrected_{file_id}_{original_filename}")
    
    if not os.path.exists(output_path):
        raise HTTPException(status_code=404, detail="修复文档不存在，请重新上传文件")
    
    base_name = os.path.splitext(original_filename)[0]
    return FileResponse(
        path=output_path,
        filename=f"{base_name}_修复版.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

@app.delete("/cleanup/{file_id}", response_description="清理临时文件")
def cleanup(file_id: str):
    if not file_id or len(file_id) != 36:
        raise HTTPException(status_code=400, detail="无效的文件ID")
    
    deleted_count = 0
    
    if os.path.exists(TEMP_DIR):
        for filename in os.listdir(TEMP_DIR):
            if file_id in filename:
                file_path = os.path.join(TEMP_DIR, filename)
                if os.path.isfile(file_path):
                    os.remove(file_path)
                    deleted_count += 1
    
    if deleted_count == 0:
        return {"message": "未找到相关文件", "deleted_count": 0}
    
    return {"message": f"成功删除 {deleted_count} 个文件", "deleted_count": deleted_count}

@app.get("/health", response_description="健康检查")
def health():
    return {
        "status": "healthy",
        "service": "Smart Text Checker API",
        "version": "1.0.0",
        "supported_formats": list(SUPPORTED_EXTENSIONS)
    }

@app.post("/analyze", response_description="分析文本（纯文本接口）")
async def analyze_text(text: str = ""):
    if not text or not text.strip():
        raise HTTPException(status_code=400, detail="文本内容不能为空")
    
    results = {
        "original_text": text,
        "errors": [],
        "corrected_text": "",
        "reference_check": [],
        "image_table_check": []
    }
    
    corrected, errors = corrector.correct_text(text)
    results["corrected_text"] = corrected
    results["errors"].extend(errors)
    
    refs = extract_references([text])
    for ref in refs:
        ref_errors = check_reference(ref)
        results["reference_check"].extend(ref_errors)
        results["errors"].extend(ref_errors)
    
    table_errors = detect_image_table_errors(text)
    results["image_table_check"].extend(table_errors)
    results["errors"].extend(table_errors)
    
    results["report"] = generate_report(results["errors"])
    
    return results

@app.post("/check/text", response_description="文本错误检测")
async def check_text_errors(file: UploadFile = File(...)):
    file_id = str(uuid.uuid4())
    os.makedirs(TEMP_DIR, exist_ok=True)
    input_path = os.path.join(TEMP_DIR, f"temp_{file_id}_{file.filename}")
    
    try:
        with open(input_path, "wb") as f:
            shutil.copyfileobj(file.file, f)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"文件保存失败: {str(e)}")
    
    try:
        text = ""
        original_html = ""
        file_ext = os.path.splitext(file.filename)[1].lower()

        if file_ext == ".pdf":
            import base64
            pdf_data = extract_pdf_structure(input_path)
            text, tables_data = _extract_pdf_body_text(input_path, pdf_data=pdf_data)
            original_html = get_pdf_html(input_path, pdf_data=pdf_data)
            images_resp = []
            for idx, img in enumerate(pdf_data.get("images", []), 1):
                img_bytes = img.get("image")
                ext = img.get("extension", "png")
                if img_bytes:
                    mime = {"png": "image/png", "jpg": "image/jpeg", "jpeg": "image/jpeg",
                            "gif": "image/gif", "bmp": "image/bmp"}.get(ext, "image/png")
                    images_resp.append({
                        "index": idx,
                        "src": f"data:{mime};base64,{base64.b64encode(img_bytes).decode('utf-8')}",
                        "caption": f"PDF图片{idx}（第{img.get('page', '?')}页）",
                        "position": f"第{img.get('page', '?')}页"
                    })
            tables_resp = []
            for idx, tb in enumerate(tables_data, 1):
                content = tb.get("content", [])
                if content:
                    tables_resp.append({
                        "index": idx,
                        "rows": content,
                        "row_count": len(content),
                        "col_count": len(content[0]) if content else 0,
                        "caption": f"PDF表格{idx}（第{tb.get('page', '?')}页）",
                        "header": content[0] if content else [],
                        "position": f"第{tb.get('page', '?')}页"
                    })
        elif file_ext in [".png", ".jpg", ".jpeg", ".bmp", ".tiff", ".gif"]:
            text = extract_text_from_image(input_path)
            # 使用相似度去重（阈值80%）
            paragraphs = text.split("\n")
            unique_paragraphs = deduplicate_by_similarity(paragraphs, similarity_threshold=0.8)
            text = "\n".join(unique_paragraphs)
        elif file_ext == ".txt":
            with open(input_path, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()
        elif file_ext == ".docx":
            doc = Document(input_path)
            text = "\n".join([para.text for para in doc.paragraphs])
            original_html = get_docx_html(input_path)
        else:
            doc = Document(input_path)
            text = "\n".join([para.text for para in doc.paragraphs])

        if file_ext != ".pdf":
            images_resp = []
            tables_resp = []

        paragraphs = text.split('\n')
        corrected_parts = []
        errors = []
        for para_text in paragraphs:
            if para_text.strip():
                corrected_para, para_errors = corrector.correct_text(para_text)
                errors.extend(para_errors)
                corrected_parts.append(corrected_para if corrected_para else para_text)
            else:
                corrected_parts.append(para_text)
        corrected = '\n'.join(corrected_parts)

        if file_ext == ".docx":
            doc_for_tables = Document(input_path)
            for tbl_idx, table in enumerate(doc_for_tables.tables, 1):
                rows = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    rows.append(row_data)
                table_data = {
                    'index': tbl_idx,
                    'rows': rows,
                    'header': rows[0] if rows else [],
                    'row_count': len(rows),
                    'col_count': len(rows[0]) if rows else 0,
                }
                table_errors = analyze_table_content(table_data, corrector)
                errors.extend(table_errors)

        if file_ext == ".pdf":
            for tbl_idx, tb in enumerate(tables_data, 1):
                content = tb.get("content", [])
                if not content:
                    continue
                table_data = {
                    'index': tbl_idx,
                    'rows': content,
                    'header': content[0] if content else [],
                    'row_count': len(content),
                    'col_count': len(content[0]) if content else 0,
                }
                table_errors = analyze_table_content(table_data, corrector)
                errors.extend(table_errors)

        details = []
        for idx, error in enumerate(errors, 1):
            pos = error.get('pos', {}).get('start', idx) if isinstance(error.get('pos'), dict) else error.get('pos', idx)
            err_text = error.get('text', '')
            suggestion = error.get('suggestion', '')
            message = error.get('message', '')
            if err_text and suggestion:
                if message:
                    details.append(f"位置{pos}：{err_text}→{suggestion}（{message}）")
                else:
                    details.append(f"位置{pos}：{err_text}→{suggestion}")
            elif err_text:
                details.append(f"位置{pos}：{err_text}")

        reference_count = sum(1 for e in errors if e.get('type') == 'reference')
        grammar_count = sum(1 for e in errors if e.get('type') == 'grammar')
        spell_count = len(errors) - reference_count - grammar_count
        
        # 为PDF生成修正文档（与Word文档一致的表格+文本修正）
        corrected_output_path = os.path.join(TEMP_DIR, f"corrected_{file_id}_{file.filename}")
        try:
            if file_ext == ".docx":
                shutil.copy(input_path, corrected_output_path)
                doc_corrected = Document(corrected_output_path)
                for para in doc_corrected.paragraphs:
                    if para.text.strip():
                        corrected_para, _ = corrector.correct_text(para.text)
                        if corrected_para and corrected_para != para.text:
                            para.text = corrected_para
                apply_table_fixes_to_docx(doc_corrected, corrector)
                doc_corrected.save(corrected_output_path)
                print(f"[INFO] 文本修正文档已生成: {corrected_output_path}")
            elif file_ext == ".pdf":
                doc_corrected = build_pdf_corrected_docx(input_path)
                if doc_corrected:
                    doc_corrected.save(corrected_output_path)
                    print(f"[INFO] PDF文本修正文档已生成: {corrected_output_path}")
                else:
                    print(f"[WARNING] PDF修正文档构建失败，使用纯文本回退")
                    fallback_paras = [t.strip() for t in text.split("\n") if t.strip()]
                    doc_corrected = Document()
                    for para_text in fallback_paras:
                        try:
                            corrected_para, _ = corrector.correct_text(para_text)
                            doc_corrected.add_paragraph(corrected_para if corrected_para else para_text)
                        except Exception:
                            doc_corrected.add_paragraph(para_text)
                    doc_corrected.save(corrected_output_path)
            elif file_ext == ".txt":
                paragraphs = [t.strip() for t in text.split("\n") if t.strip()]
                doc_corrected = Document()
                for para_text in paragraphs:
                    corrected_para, _ = corrector.correct_text(para_text)
                    doc_corrected.add_paragraph(corrected_para if corrected_para else para_text)
                doc_corrected.save(corrected_output_path)
                print(f"[INFO] TXT文本修正文档已生成: {corrected_output_path}")
        except Exception as fix_err:
            print(f"[WARNING] 生成文本修正文档失败: {fix_err}")

        return {
            "file_id": file_id,
            "original_filename": file.filename,
            "type": "text",
            "original_text": text,
            "original_html": original_html,
            "errors": errors,
            "corrected_text": corrected,
            "total_errors": len(errors),
            "images": images_resp,
            "tables": tables_resp,
            "report": {
                "错别字": spell_count,
                "语法错误": grammar_count,
                "语义错误": 0,
                "参考文献错误": reference_count,
                "图表错误": 0,
                "表格错误": 0,
                "总错误": len(errors)
            },
            "details": details
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"文本检测失败: {str(e)}")

@app.post("/check/reference", response_description="参考文献校验")
async def check_references(file: UploadFile = File(...)):
    file_id = str(uuid.uuid4())
    os.makedirs(TEMP_DIR, exist_ok=True)
    input_path = os.path.join(TEMP_DIR, f"temp_{file_id}_{file.filename}")
    
    try:
        with open(input_path, "wb") as f:
            shutil.copyfileobj(file.file, f)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"文件保存失败: {str(e)}")
    
    try:
        text = ""
        original_html = ""
        images_resp = []
        tables_resp = []
        file_ext = os.path.splitext(file.filename)[1].lower()
        
        if file_ext == ".pdf":
            import base64
            pdf_data = extract_pdf_structure(input_path)
            text, tables_data = _extract_pdf_body_text(input_path, pdf_data=pdf_data)
            original_html = get_pdf_html(input_path, pdf_data=pdf_data)
            
            for idx, img in enumerate(pdf_data.get("images", []), 1):
                img_bytes = img.get("image")
                ext = img.get("extension", "png")
                if img_bytes:
                    mime = {"png": "image/png", "jpg": "image/jpeg", "jpeg": "image/jpeg",
                            "gif": "image/gif", "bmp": "image/bmp"}.get(ext, "image/png")
                    images_resp.append({
                        "index": idx,
                        "src": f"data:{mime};base64,{base64.b64encode(img_bytes).decode('utf-8')}",
                        "caption": f"PDF图片{idx}（第{img.get('page', '?')}页）",
                        "position": f"第{img.get('page', '?')}页"
                    })
            
            for idx, tb in enumerate(tables_data, 1):
                content = tb.get("content", [])
                if content:
                    tables_resp.append({
                        "index": idx,
                        "rows": content,
                        "row_count": len(content),
                        "col_count": len(content[0]) if content else 0,
                        "header": content[0] if content else [],
                        "page": tb.get("page", 1)
                    })
        elif file_ext in [".png", ".jpg", ".jpeg", ".bmp", ".tiff", ".gif"]:
            text = extract_text_from_image(input_path)
            # 使用相似度去重（阈值80%）
            paragraphs = text.split("\n")
            unique_paragraphs = deduplicate_by_similarity(paragraphs, similarity_threshold=0.8)
            text = "\n".join(unique_paragraphs)
        elif file_ext == ".txt":
            with open(input_path, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()
        else:
            doc = Document(input_path)
            text = "\n".join([para.text for para in doc.paragraphs])
            original_html = get_docx_html(input_path)
            
            for tbl_idx, table in enumerate(doc.tables, 1):
                rows = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    rows.append(row_data)
                if rows:
                    tables_resp.append({
                        'index': tbl_idx,
                        'rows': rows,
                        'header': rows[0] if rows else [],
                        'row_count': len(rows),
                        'col_count': len(rows[0]) if rows else 0,
                    })
        
        paragraphs = text.split('\n')
        refs = extract_references(paragraphs)
        
        all_errors = []
        ref_error_count = {}
        
        for ref in refs:
            errors = check_reference(ref)
            if errors:
                ref_error_count[ref] = len(errors)
            all_errors.extend(errors)
        
        index_errors = check_index_sequence(refs)
        for err in index_errors:
            all_errors.append({
                "type": "reference",
                "level": "error",
                "text": "参考文献编号",
                "message": err
            })
        
        # 参考文献错误数 = 总错误数量
        ref_errors_count = len(all_errors)
        
        details = []
        for idx, error in enumerate(all_errors, 1):
            err_text = error.get('text', '')
            suggestion = error.get('suggestion', '')
            message = error.get('message', '')
            pos = error.get('pos', idx)

            if err_text and suggestion:
                details.append(f"【参考文献错误】位置{pos}：{err_text}→{suggestion}（{message}）")
            elif err_text:
                details.append(f"【参考文献错误】位置{pos}：{err_text}（{message}）")
            elif message:
                details.append(f"【参考文献错误】{message}")
        
        # 为docx文件生成修正后的参考文献文档
        corrected_output_path = os.path.join(TEMP_DIR, f"corrected_{file_id}_{file.filename}")
        try:
            if file_ext == ".docx":
                shutil.copy(input_path, corrected_output_path)
                doc_corrected = Document(corrected_output_path)
                
                fixed_ref_map = {}
                for ref in refs:
                    errors = check_reference(ref)
                    if errors:
                        fixed_ref_map[ref] = suggest_reference_fix(ref)
                
                if fixed_ref_map:
                    for para in doc_corrected.paragraphs:
                        para_text = para.text
                        for original_ref, fixed_ref in fixed_ref_map.items():
                            if original_ref in para_text and original_ref != fixed_ref:
                                for run in para.runs:
                                    if original_ref in run.text:
                                        run.text = run.text.replace(original_ref, fixed_ref)
                
                doc_corrected.save(corrected_output_path)
                print(f"[INFO] 参考文献修正文档已生成: {corrected_output_path}")
            else:
                # PDF: 使用build_pdf_corrected_docx生成保留图片和表格的修复文档
                if file_ext == ".pdf":
                    doc_new = build_pdf_corrected_docx(input_path, correct_text=True)
                    if doc_new:
                        # 对文档中的参考文献进行修正
                        fixed_ref_map = {}
                        for ref in refs:
                            errors = check_reference(ref)
                            if errors:
                                fixed_ref_map[ref] = suggest_reference_fix(ref)
                        
                        if fixed_ref_map:
                            for para in doc_new.paragraphs:
                                if para.text.strip():
                                    for original_ref, fixed_ref in fixed_ref_map.items():
                                        if original_ref in para.text and original_ref != fixed_ref:
                                            para.text = para.text.replace(original_ref, fixed_ref)
                        
                        doc_new.save(corrected_output_path)
                        print(f"[INFO] 参考文献修正文档已生成: {corrected_output_path}")
                    else:
                        # 备用方案：生成纯文本修复文档
                        fixed_ref_map = {}
                        for ref in refs:
                            errors = check_reference(ref)
                            if errors:
                                fixed_ref_map[ref] = suggest_reference_fix(ref)
                        
                        doc_new = Document()
                        for para_text in paragraphs:
                            if not para_text.strip():
                                continue
                            fixed_text = para_text
                            for original_ref, fixed_ref in fixed_ref_map.items():
                                if original_ref in fixed_text and original_ref != fixed_ref:
                                    fixed_text = fixed_text.replace(original_ref, fixed_ref)
                            doc_new.add_paragraph(fixed_text)
                        
                        doc_new.save(corrected_output_path)
                        print(f"[INFO] 参考文献修正文档已生成(纯文本模式): {corrected_output_path}")
                else:
                    # TXT: 生成带修正参考文献的纯文本文档
                    fixed_ref_map = {}
                    for ref in refs:
                        errors = check_reference(ref)
                        if errors:
                            fixed_ref_map[ref] = suggest_reference_fix(ref)
                    
                    doc_new = Document()
                    for para_text in paragraphs:
                        if not para_text.strip():
                            continue
                        fixed_text = para_text
                        for original_ref, fixed_ref in fixed_ref_map.items():
                            if original_ref in fixed_text and original_ref != fixed_ref:
                                fixed_text = fixed_text.replace(original_ref, fixed_ref)
                        doc_new.add_paragraph(fixed_text)
                    
                    doc_new.save(corrected_output_path)
                    print(f"[INFO] 参考文献修正文档已生成: {corrected_output_path}")
        except Exception as fix_err:
            print(f"[WARNING] 生成参考文献修正文档失败: {fix_err}")
        
        return {
            "file_id": file_id,
            "original_filename": file.filename,
            "type": "reference",
            "original_text": text,
            "original_html": original_html,
            "references": refs,
            "errors": all_errors,
            "total_errors": ref_errors_count,
            "images": images_resp,
            "tables": tables_resp,
            "report": {
                "错别字": 0,
                "语法错误": 0,
                "语义错误": 0,
                "参考文献错误": ref_errors_count,
                "图表错误": 0,
                "表格错误": 0,
                "总错误": ref_errors_count
            },
            "details": details
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"参考文献校验失败: {str(e)}")

@app.post("/check/image", response_description="图表格式检查")
async def check_images(file: UploadFile = File(...)):
    file_id = str(uuid.uuid4())
    os.makedirs(TEMP_DIR, exist_ok=True)
    input_path = os.path.join(TEMP_DIR, f"temp_{file_id}_{file.filename}")
    
    try:
        with open(input_path, "wb") as f:
            shutil.copyfileobj(file.file, f)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"文件保存失败: {str(e)}")
    
    try:
        file_ext = os.path.splitext(file.filename)[1].lower()
        
        import base64
        
        images = []
        tables = []
        image_map = {}
        
        if file_ext == ".docx":
            doc = Document(input_path)
            
            # 提取图片（从 docx 压缩包）
            try:
                with zipfile.ZipFile(input_path, 'r') as z:
                    # 提取 document.xml.rels
                    rels_path = "word/_rels/document.xml.rels"
                    if rels_path in z.namelist():
                        rels_content = z.read(rels_path).decode("utf-8")
                        for match in re.finditer(r'Id="(rId\d+)"[^>]*Target="media/([^"]+)"', rels_content):
                            rid = match.group(1)
                            img_name = match.group(2)
                            media_path = f"word/media/{img_name}"
                            if media_path in z.namelist():
                                img_data = z.read(media_path)
                                ext = img_name.split(".")[-1].lower()
                                mime_type = {
                                    "png": "image/png",
                                    "jpg": "image/jpeg",
                                    "jpeg": "image/jpeg",
                                    "gif": "image/gif",
                                    "bmp": "image/bmp",
                                    "tiff": "image/tiff",
                                    "tif": "image/tiff",
                                    "emf": "image/x-emf",
                                    "wmf": "image/x-wmf"
                                }.get(ext, "image/png")
                                b64_data = base64.b64encode(img_data).decode("utf-8")
                                image_map[rid] = f"data:{mime_type};base64,{b64_data}"
                print(f"[INFO] 共提取 {len(image_map)} 张图片资源")
            except Exception as img_err:
                print(f"[WARNING] 提取图片失败: {img_err}")
        
        image_count = 0
        table_count = 0
        
        if file_ext == ".docx":
            for element in doc.element.body:
                tag_name = element.tag.split('}')[-1]
                
                if tag_name == 'p':
                    para_texts = []
                    found_image = False
                    
                    for child in element.iter():
                        if child.tag.endswith('}t') and child.text:
                            para_texts.append(child.text)
                        elif child.tag.endswith('}drawing') or child.tag.endswith('}pict'):
                            for blip in child.iter():
                                blip_tag = blip.tag.split('}')[-1]
                                if blip_tag in ['blip', 'blipFill']:
                                    rid = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                                    if not rid:
                                        for attr_name, attr_value in blip.attrib.items():
                                            if 'embed' in attr_name:
                                                rid = attr_value
                                                break
                                    if rid:
                                        if rid in image_map:
                                            found_image = True
                                            image_count += 1
                                            images.append({
                                                "index": image_count,
                                                "src": image_map[rid],
                                                "caption": ''.join(para_texts) if para_texts else f"图片{image_count}",
                                                "position": f"段落{element.text[:20] if element.text else '未知'}..." if element.text else f"图片{image_count}"
                                            })
                                        else:
                                            found_image = True
                                            image_count += 1
                                            images.append({
                                                "index": image_count,
                                                "src": "",
                                                "caption": ''.join(para_texts) if para_texts else f"图片{image_count}",
                                                "position": f"段落{element.text[:20] if element.text else '未知'}..." if element.text else f"图片{image_count}"
                                            })
                    
                    para_text = ''.join(para_texts).strip()
                    
                    if "如图" in para_text or "下图" in para_text or "图所示" in para_text:
                        pass
                        
                elif tag_name == 'tbl':
                    table_count += 1
                    rows = []
                    row_count = 0
                    for row in element.iter():
                        if row.tag.endswith('}tr'):
                            row_count += 1
                            cells = []
                            for cell in row.iter():
                                if cell.tag.endswith('}tc'):
                                    cell_text = []
                                    for t in cell.iter():
                                        if t.tag.endswith('}t') and t.text:
                                            cell_text.append(t.text)
                                    cells.append(''.join(cell_text).strip())
                            if cells:
                                rows.append(cells)
                    
                    tables.append({
                        "index": table_count,
                        "rows": rows,
                        "row_count": row_count,
                        "col_count": len(rows[0]) if rows else 0,
                        "caption": f"表格{table_count}",
                        "header": rows[0] if rows else [],
                        "position": f"第{table_count}个表格"
                    })

        elif file_ext == ".pdf":
            pdf_data = extract_pdf_structure(input_path)
            for idx, img in enumerate(pdf_data.get("images", []), 1):
                img_bytes = img.get("image")
                ext = img.get("extension", "png")
                if img_bytes:
                    mime = {"png": "image/png", "jpg": "image/jpeg", "jpeg": "image/jpeg",
                            "gif": "image/gif", "bmp": "image/bmp"}.get(ext, "image/png")
                    images.append({
                        "index": idx,
                        "src": f"data:{mime};base64,{base64.b64encode(img_bytes).decode('utf-8')}",
                        "caption": f"PDF图片{idx}（第{img.get('page', '?')}页）",
                        "position": f"第{img.get('page', '?')}页"
                    })
            image_count = len(images)

            for idx, tb in enumerate(pdf_data.get("tables", []), 1):
                content = tb.get("content", [])
                if content:
                    table_count += 1
                    tables.append({
                        "index": idx,
                        "rows": content,
                        "row_count": len(content),
                        "col_count": len(content[0]) if content else 0,
                        "caption": f"PDF表格{idx}（第{tb.get('page', '?')}页）",
                        "header": content[0] if content else [],
                        "position": f"第{tb.get('page', '?')}页"
                    })
        
        errors = []
        details = []
        
        # 使用新的表格分析函数进行详细检测
        for table in tables:
            table_errors = analyze_table_structure(table)
            errors.extend(table_errors)

            for err in table_errors:
                level_tag = {
                    "error": "【表格错误】",
                    "warning": "【表格警告】",
                    "info": "【表格提示】"
                }.get(err.get('level'), "【表格问题】")
                message = err.get('message', '')
                text = err.get('text', '')
                suggestion = err.get('suggestion', '')
                corrections = err.get('corrections', [])

                if corrections and len(corrections) > 0:
                    correction = corrections[0]
                    if suggestion:
                        details.append(f"{level_tag}{text}→{correction}（{suggestion}）")
                    else:
                        details.append(f"{level_tag}{text}→{correction}（{message}）")
                else:
                    details.append(f"{level_tag}{message}（{suggestion}）")
            
            # 表格内文字内容纠错（与文本错误检测一致）
            table_content_errors = analyze_table_content(table, corrector)
            errors.extend(table_content_errors)
            for e in table_content_errors:
                err_text = e.get('text', '')
                suggestion = e.get('suggestion', '')
                message = e.get('message', '')
                pos = e.get('pos', '')
                if err_text and suggestion:
                    details.append(f"【表格内容错误】位置{pos}：{err_text}→{suggestion}（{message}）")
                elif err_text:
                    details.append(f"【表格内容错误】位置{pos}：{err_text}（{message}）")
                else:
                    details.append(f"【表格内容错误】{message}")
        
        # 检测图片问题
        for img in images:
            # 检测图片是否有标题
            if not img['caption'] or not img['caption'].strip():
                errors.append({
                    "type": "image",
                    "level": "warning",
                    "pos": img['index'],
                    "text": f"图片{img['index']}",
                    "suggestion": "为图片添加规范的图题（如：图1 XXX）",
                    "message": f"图片{img['index']}缺少标题，建议添加规范图题"
                })
                details.append(f"【图表警告】图片{img['index']}缺少标题，建议添加规范图题")
            else:
                # 检测图片编号是否规范
                if not re.match(r'^图\d+\s', img['caption']):
                    errors.append({
                        "type": "image",
                        "level": "warning",
                        "pos": img['index'],
                        "text": img['caption'],
                        "suggestion": f"添加规范图片编号'图{img['index']}'作为标题前缀",
                        "message": f"图片{img['index']}缺少规范编号，建议添加'图{img['index']}'前缀"
                    })
                    details.append(f"【图表警告】图片{img['index']}缺少规范编号")
        
        if image_count == 0:
            errors.append({
                "type": "image",
                "level": "info",
                "pos": None,
                "text": "文档图片",
                "suggestion": "文档中未发现图片",
                "message": "文档中未提取到任何图片"
            })
            details.append(f"【图表错误】文档中未提取到任何图片")
        
        if table_count == 0:
            errors.append({
                "type": "table",
                "level": "info",
                "pos": None,
                "text": "文档表格",
                "suggestion": "文档中未发现表格",
                "message": "文档中未提取到任何表格"
            })
            details.append(f"【表格错误】文档中未提取到任何表格")
        
        image_errors = [e for e in errors if e.get('type') == 'image']
        table_errors = [e for e in errors if e.get('type') == 'table']
        spell_errors = [e for e in errors if e.get('type') == 'spell']
        
        # 生成修正文档（表格内容纠错 + 文本内容）
        corrected_output_path = os.path.join(TEMP_DIR, f"corrected_{file_id}_{file.filename}")
        try:
            if file_ext == ".docx":
                shutil.copy(input_path, corrected_output_path)
                doc_corrected = Document(corrected_output_path)
                apply_table_fixes_to_docx(doc_corrected, corrector)
                doc_corrected.save(corrected_output_path)
                print(f"[INFO] 图表表格修正文档已生成: {corrected_output_path}")
            elif file_ext == ".pdf":
                doc_corrected = build_pdf_corrected_docx(input_path, correct_text=False)
                if doc_corrected:
                    doc_corrected.save(corrected_output_path)
                    print(f"[INFO] 图表表格修正文档已生成: {corrected_output_path}")
                else:
                    print(f"[WARNING] PDF修正文档构建失败，跳过修正文档生成")
        except Exception as fix_err:
            print(f"[WARNING] 生成修正文档失败: {fix_err}")
        
        # 为PDF生成 original_html 和 original_text
        original_text = ""
        original_html = ""
        if file_ext == ".pdf":
            pdf_img_data = extract_pdf_structure(input_path)
            original_html = get_pdf_html(input_path, pdf_data=pdf_img_data)
            original_text, _ = _extract_pdf_body_text(input_path, pdf_data=pdf_img_data)
        elif file_ext == ".docx":
            original_html = get_docx_html(input_path)
            original_text = "\n".join([para.text for para in doc.paragraphs])
        
        return {
            "file_id": file_id,
            "original_filename": file.filename,
            "type": "image",
            "images": images,
            "tables": tables,
            "original_html": original_html,
            "original_text": original_text,
            "image_count": image_count,
            "table_count": table_count,
            "errors": errors,
            "total_errors": len(errors),
            "report": {
                "错别字": len(spell_errors),
                "语法错误": 0,
                "语义错误": 0,
                "参考文献错误": 0,
                "图表错误": len(image_errors),
                "表格错误": len(table_errors),
                "总错误": len(errors)
            },
            "details": details
        }
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"图表检查失败: {str(e)}")

@app.post("/check/table", response_description="表格格式检查")
async def check_tables(file: UploadFile = File(...)):
    file_id = str(uuid.uuid4())
    os.makedirs(TEMP_DIR, exist_ok=True)
    input_path = os.path.join(TEMP_DIR, f"temp_{file_id}_{file.filename}")
    
    try:
        with open(input_path, "wb") as f:
            shutil.copyfileobj(file.file, f)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"文件保存失败: {str(e)}")
    
    try:
        text = ""
        file_ext = os.path.splitext(file.filename)[1].lower()
        
        if file_ext == ".pdf":
            text, _ = _extract_pdf_body_text(input_path)
        elif file_ext in [".png", ".jpg", ".jpeg", ".bmp", ".tiff", ".gif"]:
            text = extract_text_from_image(input_path)
            # 使用相似度去重（阈值80%）
            paragraphs = text.split("\n")
            unique_paragraphs = deduplicate_by_similarity(paragraphs, similarity_threshold=0.8)
            text = "\n".join(unique_paragraphs)
        elif file_ext == ".txt":
            with open(input_path, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()
        else:
            doc = Document(input_path)
            text = "\n".join([para.text for para in doc.paragraphs])
        
        errors = detect_image_table_errors(text)
        table_errors = [e for e in errors if '表' in e.get('message', '')]
        
        details = []
        for error in table_errors:
            message = error.get('message', '')
            if message:
                details.append(f"【表格错误】{message}")
        
        return {
            "file_id": file_id,
            "original_filename": file.filename,
            "type": "table",
            "original_text": text,
            "errors": table_errors,
            "total_errors": len(table_errors),
            "report": {
                "错别字": 0,
                "语法错误": 0,
                "语义错误": 0,
                "参考文献错误": 0,
                "图表错误": 0,
                "表格错误": len(table_errors),
                "总错误": len(table_errors)
            },
            "details": details
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"表格检查失败: {str(e)}")

@app.post("/file/info", response_description="获取文件信息")
async def get_file_info(file: UploadFile = File(...)):
    file_id = str(uuid.uuid4())
    os.makedirs(TEMP_DIR, exist_ok=True)
    input_path = os.path.join(TEMP_DIR, f"temp_{file_id}_{file.filename}")
    
    try:
        with open(input_path, "wb") as f:
            shutil.copyfileobj(file.file, f)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"文件保存失败: {str(e)}")
    
    file_size = os.path.getsize(input_path)
    file_ext = os.path.splitext(file.filename)[1].lower()
    
    type_map = {
        ".docx": "Word文档",
        ".pdf": "PDF文档",
        ".txt": "文本文件",
        ".png": "PNG图片",
        ".jpg": "JPEG图片",
        ".jpeg": "JPEG图片",
        ".bmp": "BMP图片",
        ".tiff": "TIFF图片",
        ".gif": "GIF图片"
    }
    
    file_type = type_map.get(file_ext, "未知类型")
    
    return {
        "filename": file.filename,
        "size": file_size,
        "size_formatted": f"{file_size / 1024:.2f} KB" if file_size > 1024 else f"{file_size} B",
        "extension": file_ext,
        "type": file_type
    }

@app.post("/fix/table", response_description="修正表格问题")
async def fix_table(
    file: UploadFile = File(...),
    table_index: int = None,
    new_caption: str = None,
    new_header: bool = False,
    add_unit: bool = False,
    add_note: bool = False
):
    """
    修正Word文档中的表格问题
    
    参数:
    - file: Word文档
    - table_index: 要修正的表格索引（从1开始）
    - new_caption: 新的表格标题（可选）
    - new_header: 是否添加单位到表头（可选）
    - add_unit: 是否为数值列添加单位（可选）
    - add_note: 是否添加表注（可选）
    """
    file_id = str(uuid.uuid4())
    os.makedirs(TEMP_DIR, exist_ok=True)
    input_path = os.path.join(TEMP_DIR, f"temp_{file_id}_{file.filename}")
    output_path = os.path.join(TEMP_DIR, f"fixed_{file_id}_{file.filename}")
    
    try:
        with open(input_path, "wb") as f:
            shutil.copyfileobj(file.file, f)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"文件保存失败: {str(e)}")
    
    try:
        doc = Document(input_path)
        
        tables = doc.tables
        if table_index is None or table_index < 1 or table_index > len(tables):
            raise HTTPException(status_code=400, detail=f"无效的表格索引，有效范围为1到{len(tables)}")
        
        table = tables[table_index - 1]
        modifications = []
        
        if new_caption:
            # 在表格前查找并更新标题段落
            # 简化处理：如果有新标题，直接在文档中查找包含"表格"的段落并替换
            for para in doc.paragraphs:
                if f"表格{table_index}" in para.text or (table_index > 0 and f"表{table_index}" in para.text):
                    # 保留原有样式，只修改文本
                    for run in para.runs:
                        if "表格" in run.text or "表" in run.text:
                            run.text = new_caption
                            modifications.append(f"标题已更新为: {new_caption}")
                            break
                    else:
                        para.text = new_caption
                        modifications.append(f"标题已更新为: {new_caption}")
                    break
        
        if new_header and table.rows:
            # 在表头添加单位
            header_row = table.rows[0]
            for i, cell in enumerate(header_row.cells):
                cell_text = cell.text.strip()
                # 检查是否已经有单位（括号、%或单位后缀）
                has_unit = any(u in cell_text for u in ['%', '（', '）', '(', ')', '(%)', '(s)', '(单位)', '（%）', '（s）', '（单位）'])
                if cell_text and not has_unit:
                    # 检查是否是数值列
                    data_row = table.rows[1] if len(table.rows) > 1 else None
                    if data_row and i < len(data_row.cells):
                        cell_data = data_row.cells[i].text.strip()
                        # 如果数据行包含数字，添加单位
                        if cell_data and any(c.isdigit() for c in cell_data):
                            # 推断单位
                            unit = ''
                            if '率' in cell_text or '准确' in cell_text or '召回' in cell_text:
                                unit = '(%)'
                            elif 'F1' in cell_text:
                                unit = '(%)'
                            elif '时间' in cell_text or '耗时' in cell_text:
                                unit = '(s)'
                            
                            if unit and unit not in cell_text:
                                new_text = f"{cell_text}{unit}"
                                for run in cell.paragraphs[0].runs if cell.paragraphs else []:
                                    run.text = new_text
                                    modifications.append(f"表头'{cell_text}'已添加单位{unit}")
                                    break
                                else:
                                    cell.text = new_text
                                    modifications.append(f"表头'{cell_text}'已添加单位{unit}")
        
        if add_note:
            # 在表格后添加表注
            last_para = None
            for para in doc.paragraphs:
                last_para = para
            
            if last_para is not None:
                note_text = "注：实验数据来源于XX，数据处理方法为XX。"
                new_para = doc.add_paragraph(note_text)
                modifications.append(f"已添加表注: {note_text}")
        
        doc.save(output_path)
        
        return {
            "success": True,
            "file_id": file_id,
            "filename": f"fixed_{file.filename}",
            "output_path": output_path,
            "modifications": modifications,
            "message": "表格修正完成"
        }
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"表格修正失败: {str(e)}")

@app.post("/edit/table", response_description="编辑表格内容")
async def edit_table(
    file: UploadFile = File(...),
    table_index: int = None,
    headers: str = None,
    rows: str = None
):
    """
    编辑Word文档中的表格内容
    
    参数:
    - file: Word文档
    - table_index: 要编辑的表格索引（从1开始）
    - headers: 表头内容（JSON数组字符串）
    - rows: 表格行数据（JSON数组字符串）
    """
    file_id = str(uuid.uuid4())
    os.makedirs(TEMP_DIR, exist_ok=True)
    input_path = os.path.join(TEMP_DIR, f"temp_{file_id}_{file.filename}")
    output_path = os.path.join(TEMP_DIR, f"edited_{file_id}_{file.filename}")
    
    try:
        with open(input_path, "wb") as f:
            shutil.copyfileobj(file.file, f)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"文件保存失败: {str(e)}")
    
    try:
        doc = Document(input_path)
        
        tables = doc.tables
        print(f"[DEBUG] /edit/table - 接收到表格索引: {table_index}, 文档中表格总数: {len(tables)}")
        if table_index is None or table_index < 1 or table_index > len(tables):
            raise HTTPException(status_code=400, detail=f"无效的表格索引 {table_index}，有效范围为1到{len(tables)}")
        
        table = tables[table_index - 1]
        modifications = []
        
        if headers:
            try:
                headers_list = json.loads(headers)
                header_row = table.rows[0]
                for i, header in enumerate(headers_list):
                    if i < len(header_row.cells):
                        old_text = header_row.cells[i].text.strip()
                        if old_text != header.strip():
                            header_row.cells[i].text = header
                            modifications.append(f"表头第{i+1}列: '{old_text}' -> '{header}'")
            except json.JSONDecodeError:
                raise HTTPException(status_code=400, detail="headers参数格式错误")
        
        if rows:
            try:
                rows_list = json.loads(rows)
                for row_idx, row_data in enumerate(rows_list):
                    table_row_idx = row_idx + 1
                    if table_row_idx < len(table.rows):
                        table_row = table.rows[table_row_idx]
                        for col_idx, cell_data in enumerate(row_data):
                            if col_idx < len(table_row.cells):
                                old_text = table_row.cells[col_idx].text.strip()
                                if old_text != cell_data.strip():
                                    table_row.cells[col_idx].text = cell_data
                                    modifications.append(f"第{table_row_idx+1}行第{col_idx+1}列: '{old_text}' -> '{cell_data}'")
            except json.JSONDecodeError:
                raise HTTPException(status_code=400, detail="rows参数格式错误")
        
        doc.save(output_path)
        
        return {
            "success": True,
            "file_id": file_id,
            "filename": f"edited_{file.filename}",
            "output_path": output_path,
            "modifications": modifications,
            "message": "表格编辑完成"
        }
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"表格编辑失败: {str(e)}")

@app.post("/download/fixed", response_description="下载修正后的文件")
async def download_fixed(request: dict):
    """
    下载修正后的文件
    """
    file_id = request.get('file_id')
    filename = request.get('filename')
    
    if not file_id or not filename:
        raise HTTPException(status_code=400, detail="缺少file_id或filename参数")
    
    # 查找修正后的文件
    temp_files = os.listdir(TEMP_DIR) if os.path.exists(TEMP_DIR) else []
    base_name = os.path.splitext(filename)[0]
    
    # 1. 优先查找 fixed_ 开头的文件（由 /fix/all 生成）
    # 必须同时匹配 file_id 和 filename
    for f in temp_files:
        if f.startswith("fixed_") and file_id in f and filename in f:
            file_path = os.path.join(TEMP_DIR, f)
            return FileResponse(
                file_path,
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                filename=f"{base_name}_修正版.docx"
            )
    
    # 2. 查找 corrected_ 开头的文件（由 /upload 生成）
    # 必须匹配 file_id
    for f in temp_files:
        if f.startswith("corrected_") and file_id in f:
            f_base = os.path.splitext(f)[0]
            # 匹配：文件名包含 base_name 或 original filename
            if filename in f or base_name in f_base:
                file_path = os.path.join(TEMP_DIR, f)
                return FileResponse(
                    file_path,
                    media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    filename=f"{base_name}_修复版.docx"
                )
    
    # 没有找到匹配的文件
    raise HTTPException(status_code=404, detail="文件不存在或已过期，请重新生成修复文档")

@app.post("/fix/all", response_description="一键修正所有错误")
async def fix_all_errors(file: UploadFile = File(...), error_indices: str = Form(None)):
    """
    一键修正所有错误
    :param file: 原始文档文件
    :param error_indices: 需要修正的错误索引列表（逗号分隔，可选，不传则修正全部）
    """
    file_id = str(uuid.uuid4())
    os.makedirs(TEMP_DIR, exist_ok=True)
    input_path = os.path.join(TEMP_DIR, f"temp_{file_id}_{file.filename}")
    
    try:
        with open(input_path, "wb") as f:
            shutil.copyfileobj(file.file, f)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"文件保存失败: {str(e)}")
    
    try:
        file_ext = os.path.splitext(file.filename)[1].lower()
        
        if file_ext == ".pdf":
            output_path = os.path.join(TEMP_DIR, f"fixed_{file_id}_{file.filename}.docx")
            doc_corrected = build_pdf_corrected_docx(input_path)
            if not doc_corrected:
                raise HTTPException(status_code=500, detail="PDF修正文档生成失败")
            doc_corrected.save(output_path)
            print(f"[INFO] PDF修正文档已生成: {output_path}")
            return {
                "file_id": file_id,
                "filename": f"fixed_{file_id}_{file.filename}.docx",
                "original_filename": file.filename,
                "output_path": output_path,
                "corrected_count": 0,
                "modifications": ["PDF已转换为修正版Word文档"],
                "message": "PDF修正版文档生成完成"
            }
        
        if file_ext != ".docx":
            raise HTTPException(status_code=400, detail="批量修正仅支持Word和PDF文档格式")
        
        output_path = os.path.join(TEMP_DIR, f"fixed_{file_id}_{file.filename}")
        shutil.copy(input_path, output_path)
        
        doc = Document(output_path)
        
        selected_indices = None
        if error_indices:
            selected_indices = [int(i.strip()) for i in error_indices.split(",") if i.strip().isdigit()]
        
        modifications = []
        error_counter = 0
        
        for para in doc.paragraphs:
            if para.text.strip():
                corrected_para, para_errors = corrector.correct_text(para.text)
                
                if para_errors:
                    modified_text = para.text
                    for error in para_errors:
                        err_text = error.get('text', '')
                        suggestion = error.get('suggestion', '')
                        if err_text and suggestion:
                            if selected_indices is None or error_counter in selected_indices:
                                modified_text = modified_text.replace(err_text, suggestion)
                                modifications.append(f"修正: '{err_text}' → '{suggestion}'")
                            error_counter += 1
                    para.text = modified_text
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if para.text.strip():
                            corrected_para, para_errors = corrector.correct_text(para.text)
                            
                            if para_errors:
                                modified_text = para.text
                                for error in para_errors:
                                    err_text = error.get('text', '')
                                    suggestion = error.get('suggestion', '')
                                    if err_text and suggestion:
                                        if selected_indices is None or error_counter in selected_indices:
                                            modified_text = modified_text.replace(err_text, suggestion)
                                            modifications.append(f"修正: '{err_text}' → '{suggestion}'")
                                        error_counter += 1
                                para.text = modified_text
        
        # 应用表格内容修正（单位补充等）
        apply_table_fixes_to_docx(doc, corrector)
        
        doc.save(output_path)
        
        return {
            "file_id": file_id,
            "filename": f"fixed_{file_id}_{file.filename}",
            "original_filename": file.filename,
            "output_path": output_path,
            "modifications": modifications,
            "corrected_count": len(modifications),
            "message": "批量修正完成"
        }
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"批量修正失败: {str(e)}")

print("[INFO] API服务启动完成")
    