#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
智能文本检测与格式批注系统 - 本地运行入口
融合NLP文本检测与计算机视觉技术，开发论文/报告智能校验工具

核心功能：
1. 错别字识别（支持双模型结构：macbert_finetuned + pycorrector）
2. 语法错误检测
3. 参考文献格式校验（按 GB/T 7714 规范）
4. 图片编号不规范检测
5. 表格内容检测
6. 自动生成批注与一键修正建议

支持的纠错模式：
- single: 单模型模式（仅 macbert_finetuned，不带 pycorrector）
- dual: 双模型模式（macbert_finetuned + pycorrector，深度学习+规则引擎）
"""
import re
import os
import sys
import zipfile
import torch
import argparse
import logging
from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 添加当前目录到路径
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

# 导入模块
from reporter import generate_report, generate_detail_report
from nlp_corrector import TextCorrector
from pdf_processor import extract_text_from_pdf
try:
    from image_processor import extract_text_from_image, analyze_image_quality
    HAS_IMAGE_PROCESSOR = True
except ImportError:
    from image_ocr import extract_text_from_image
    HAS_IMAGE_PROCESSOR = False
from reference_checker import (
    extract_references,
    check_reference,
    suggest_reference_fix,
    check_index_sequence
)
from image_table_checker import detect_image_table_errors
from grammar_checker import check_syntax, analyze_sentence_structure, check_academic_expressions, check_semantic_errors

# ===== 日志配置 =====
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.StreamHandler(),
        logging.FileHandler('checker.log', encoding='utf-8')
    ]
)
logger = logging.getLogger(__name__)

# 修复Windows命令行中文显示问题
import io
import sys
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8')

# ===== 配置常量 =====
BASE_DIR = os.path.dirname(os.path.dirname(__file__))

# 模型配置（使用 macbert_finetuned + pycorrector 双模型）
MODEL_PATH_V3 = os.path.join(BASE_DIR, "models", "macbert_finetuned")

# 默认使用双模型模式（macbert_finetuned + pycorrector）
DEFAULT_MODE = "dual"  # "single", "dual"

# ===== 颜色定义 =====
COLOR_ERROR = RGBColor(220, 53, 69)      # 红色 - 错误
COLOR_WARNING = RGBColor(255, 193, 7)    # 橙色 - 警告
COLOR_SUCCESS = RGBColor(40, 167, 69)    # 绿色 - 正确
COLOR_INFO = RGBColor(41, 128, 185)      # 蓝色 - 信息
COLOR_TEXT = RGBColor(51, 51, 51)        # 黑色 - 正文

class AcademicDocumentChecker:
    """学术文档校验器（支持双模型结构）"""
    
    def __init__(self, mode=DEFAULT_MODE):
        """初始化校验器
        
        Args:
            mode: 纠错模式，"single"（单模型+pycorrector）或 "dual"（双模型）
        """
        self.device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
        logger.info(f"使用设备: {self.device}")
        
        self.corrector = None
        self.mode = mode.lower()
        
    def load_model(self):
        """加载NLP模型（支持单模型、双模型模式）"""
        logger.info(f"[INFO] 加载模型中... (模式: {self.mode})")
        try:
            if self.mode == "dual":
                # 双模型模式：macbert_finetuned + pycorrector（深度学习 + 规则引擎）
                logger.info("[INFO] 加载双模型: macbert_finetuned + pycorrector")
                self.corrector = TextCorrector(MODEL_PATH_V3, self.device)
            else:
                # 单模型模式：仅使用macbert_finetuned（不带pycorrector）
                logger.info("[INFO] 加载单模型: macbert_finetuned")
                self.corrector = TextCorrector(MODEL_PATH_V3, self.device, use_pycorrector=False)
            
            logger.info("[OK] 模型加载完成")
            return True
        except Exception as e:
            logger.error(f"[ERROR] 模型加载失败: {e}")
            return False
    
    def extract_images_from_docx(self, docx_path, output_dir="temp_images"):
        """提取Word文档中的图片"""
        os.makedirs(output_dir, exist_ok=True)
        image_paths = []
        
        try:
            with zipfile.ZipFile(docx_path, 'r') as zip_ref:
                for file in zip_ref.namelist():
                    if file.startswith("word/media/"):
                        zip_ref.extract(file, output_dir)
            
            media_path = os.path.join(output_dir, "word", "media")
            if os.path.exists(media_path):
                for img in os.listdir(media_path):
                    image_paths.append(os.path.join(media_path, img))
            logger.info(f"📸 提取到 {len(image_paths)} 张图片")
        except Exception as e:
            logger.error(f"❌ 提取图片失败: {e}")
        
        return image_paths
    
    def check_image_numbering(self, paragraphs):
        """检查图片编号是否规范"""
        errors = []
        image_pattern = re.compile(r'图\s*(\d+)\s*[-—–]\s*(\d+)|图\s*(\d+)', re.IGNORECASE)
        numbers = []
        
        for idx, p in enumerate(paragraphs):
            match = image_pattern.search(p)
            if match:
                if match.group(1):  # 图X-Y 格式
                    num = f"{match.group(1)}-{match.group(2)}"
                else:  # 图X 格式
                    num = match.group(3)
                numbers.append((idx, num))
        
        # 检查编号连续性
        if numbers:
            int_nums = []
            for idx, num in numbers:
                if '-' in num:
                    parts = num.split('-')
                    if len(parts) == 2 and parts[0].isdigit():
                        int_nums.append((idx, int(parts[0])))
                elif num.isdigit():
                    int_nums.append((idx, int(num)))
            
            int_nums.sort(key=lambda x: x[1])
            for i, (idx, num) in enumerate(int_nums):
                expected = i + 1
                if num != expected:
                    errors.append({
                        "type": "image",
                        "level": "error",
                        "pos": idx,
                        "text": f"图{num}",
                        "suggestion": f"图{expected}",
                        "message": f"图片编号不连续，应为图{expected}"
                    })
        
        return errors
    
    def check_table_numbering(self, doc):
        """检查表格编号是否规范"""
        errors = []
        table_nums = []
        
        for idx, table in enumerate(doc.tables):
            # 检查表格标题（通常在表格前）
            if idx > 0:
                prev_para = doc.paragraphs[idx]
                match = re.search(r'表\s*(\d+)', prev_para.text, re.IGNORECASE)
                if match:
                    table_nums.append((idx, int(match.group(1))))
        
        # 检查编号连续性
        table_nums.sort(key=lambda x: x[1])
        for i, (idx, num) in enumerate(table_nums):
            expected = i + 1
            if num != expected:
                errors.append({
                    "type": "table",
                    "level": "error",
                    "pos": idx,
                    "text": f"表{num}",
                    "suggestion": f"表{expected}",
                    "message": f"表格编号不连续，应为表{expected}"
                })
        
        return errors
    
    def process_document(self, input_file, output_file):
        """处理文档主流程"""
        all_errors = []
        
        # ===== 1. 读取文件 =====
        logger.info(f"📄 读取文件: {input_file}")
        try:
            if input_file.endswith(".docx"):
                doc = Document(input_file)
                paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
                is_docx = True
                logger.info(f"✅ 读取成功，共 {len(paragraphs)} 段")
                
            elif input_file.endswith(".pdf"):
                from pdf_processor import extract_pdf_structure
                pdf_data = extract_pdf_structure(input_file)
                doc = Document()
                is_docx = False
                logger.info(f"✅ PDF解析成功，共 {len(pdf_data['text'])} 页文本, {len(pdf_data['images'])} 张图片, {len(pdf_data['tables'])} 个表格")
                
            elif input_file.endswith((".png", ".jpg", ".jpeg")):
                text = extract_text_from_image(input_file)
                paragraphs = [line.strip() for line in text.split("\n") if line.strip()]
                doc = Document()
                is_docx = False
                logger.info(f"✅ 图片OCR成功，共 {len(paragraphs)} 段")
                
            else:
                raise ValueError("仅支持 docx / pdf / image 格式")
                
        except Exception as e:
            logger.error(f"❌ 文件读取失败: {e}")
            return None, []
        
        # ===== 2. 正文处理与纠错 =====
        logger.info("📝 开始文本纠错...")
        try:
            if is_docx:
                # 处理现有段落 - 直接修改原段落
                for para in doc.paragraphs:
                    text = para.text.strip()
                    if not text:
                        continue
                    
                    corrected, errors = self.corrector.correct_text(text)
                    
                    if errors:
                        # 清空原段落并替换为纠错内容
                        self._build_correction_paragraph(para, text, corrected, errors, keep_images=True)
                        all_errors.extend(errors)
            else:
                # PDF处理：新建文档，保留结构
                if input_file.endswith(".pdf"):
                    # 处理PDF文本内容（按页）
                    for page_data in pdf_data["text"]:
                        page_num = page_data["page"]
                        page_text = page_data["content"]
                        
                        # 按段落分割
                        page_paragraphs = [p.strip() for p in page_text.split("\n\n") if p.strip()]
                        
                        for text in page_paragraphs:
                            if not text:
                                continue
                            
                            corrected, errors = self.corrector.correct_text(text)
                            
                            if errors:
                                # 直接创建纠错段落
                                para = doc.add_paragraph()
                                run = para.add_run(f"【第{page_num}页】")
                                run.font.color.rgb = COLOR_WARNING
                                run.font.size = Pt(10)
                                run.bold = True
                                self._build_correction_paragraph(para, text, corrected, errors)
                                all_errors.extend(errors)
                            else:
                                para = doc.add_paragraph()
                                run = para.add_run(f"【第{page_num}页】{text}")
                                run.font.color.rgb = COLOR_TEXT
                                run.font.size = Pt(12)
                
                # 图片/OCR处理：新建文档
                else:
                    for text in paragraphs:
                        corrected, errors = self.corrector.correct_text(text)
                        
                        if errors:
                            # 直接创建纠错段落
                            para = doc.add_paragraph()
                            self._build_correction_paragraph(para, text, corrected, errors)
                            all_errors.extend(errors)
                        else:
                            para = doc.add_paragraph()
                            run = para.add_run(text)
                            run.font.color.rgb = COLOR_TEXT
                            run.font.size = Pt(12)
            
            logger.info(f"✅ 文本纠错完成，发现 {len([e for e in all_errors if e.get('type') == 'spell'])} 个错别字")
            
            # ===== 语法与语义检测 =====
            logger.info("🔍 开始语法检测...")
            
            # 先收集所有需要处理的段落（避免在循环中修改列表）
            paragraphs_to_process = list(doc.paragraphs)
            
            for i, para in enumerate(paragraphs_to_process):
                text = para.text.strip()
                if not text:
                    continue
                
                # 跳过已处理的段落（包含标记的段落）
                if any(marker in text for marker in ['【原文】', '【建议】', '【图片识别】', '【第', '【表格编号', '【图片编号', '【参考文献']):
                    continue
                
                # 收集该段落的所有语法错误
                paragraph_errors = []
                
                # 检测语法问题
                syntax_errors = check_syntax(text)
                paragraph_errors.extend(syntax_errors)
                all_errors.extend(syntax_errors)
                
                # 分析句子结构
                structure_suggestions = analyze_sentence_structure(text)
                paragraph_errors.extend(structure_suggestions)
                all_errors.extend(structure_suggestions)
                
                # 检查学术表达
                academic_errors = check_academic_expressions(text)
                paragraph_errors.extend(academic_errors)
                all_errors.extend(academic_errors)
                
                # 检测语义错误
                semantic_errors = check_semantic_errors(text)
                paragraph_errors.extend(semantic_errors)
                all_errors.extend(semantic_errors)
                
                # 检测图片表格引用规范
                image_table_errors = detect_image_table_errors(text)
                paragraph_errors.extend(image_table_errors)
                all_errors.extend(image_table_errors)
                
                # 在该段落下方直接添加语法检测结果（合并显示）
                if paragraph_errors:
                    # 添加分隔线
                    line_para = doc.add_paragraph()
                    run = line_para.add_run("=" * 60)
                    run.font.color.rgb = COLOR_INFO
                    
                    # 显示原文
                    orig_para = doc.add_paragraph()
                    run = orig_para.add_run("【原句】")
                    run.font.color.rgb = RGBColor(0, 0, 0)  # 黑色
                    run.bold = True
                    run = orig_para.add_run(text)
                    run.font.color.rgb = RGBColor(0, 0, 0)  # 黑色
                    
                    # 显示所有错误（合并显示）
                    error_para = doc.add_paragraph()
                    run = error_para.add_run("【问题】")
                    run.font.color.rgb = RGBColor(192, 0, 0)  # 深红色
                    run.bold = True
                    
                    for i, error in enumerate(paragraph_errors, 1):
                        if i > 1:
                            run = error_para.add_run("；")
                        run = error_para.add_run(f"{i}.{error['message']}")
                        run.font.color.rgb = RGBColor(192, 0, 0)  # 深红色
                    
                    # 显示所有建议（合并显示）
                    suggest_para = doc.add_paragraph()
                    run = suggest_para.add_run("【建议】")
                    run.font.color.rgb = RGBColor(0, 102, 0)  # 深绿色
                    run.bold = True
                    
                    for i, error in enumerate(paragraph_errors, 1):
                        if i > 1:
                            run = suggest_para.add_run("；")
                        run = suggest_para.add_run(f"{i}.{error['suggestion']}")
                        run.font.color.rgb = RGBColor(0, 102, 0)  # 深绿色
                    
                    # 显示修改后的句子（根据建议自动修复）
                    corrected_text = text
                    import re as regex_module
                    for error in paragraph_errors:
                        # 解析"改为「XXX」"格式的建议
                        suggest_match = regex_module.search(r'改为「([^」]+)」', error['suggestion'])
                        if suggest_match:
                            # 获取建议替换的内容
                            replacement = suggest_match.group(1)
                            # 获取匹配到的原始文本模式
                            error_pattern = error.get('pattern', '')
                            if error_pattern:
                                # 使用正则表达式替换
                                corrected_text = regex_module.sub(error_pattern, replacement, corrected_text)
                        elif error.get('fix'):
                            corrected_text = error['fix'](corrected_text)
                    
                    fix_para = doc.add_paragraph()
                    run = fix_para.add_run("【修改后】")
                    run.font.color.rgb = RGBColor(0, 102, 0)  # 深绿色
                    run.bold = True
                    run = fix_para.add_run(corrected_text.strip())
                    run.font.color.rgb = RGBColor(0, 102, 0)  # 深绿色
                    
                    # 添加分隔线
                    line_para = doc.add_paragraph()
                    run = line_para.add_run("=" * 60)
                    run.font.color.rgb = COLOR_INFO
            
            grammar_count = len([e for e in all_errors if e.get('type') == 'grammar'])
            semantic_count = len([e for e in all_errors if e.get('type') == 'semantic'])
            logger.info(f"✅ 语法检测完成，发现 {grammar_count} 个语法问题，{semantic_count} 个语义问题")
        
        except Exception as e:
            logger.error(f"❌ 正文处理失败: {e}")
        
        # ===== 3. 表格处理 =====
        if is_docx:
            logger.info("📊 处理表格...")
            try:
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            text = cell.text.strip()
                            if not text:
                                continue
                            
                            corrected, errors = self.corrector.correct_text(text)
                            
                            if errors:
                                cell.text = f"【原文】{text}\n【建议】{corrected}"
                                all_errors.extend(errors)
                logger.info("✅ 表格处理完成")
            except Exception as e:
                logger.error(f"❌ 表格处理失败: {e}")
        
        # PDF表格处理
        if input_file.endswith(".pdf"):
            logger.info("📊 处理PDF表格...")
            try:
                from pdf_processor import extract_pdf_structure
                if pdf_data and pdf_data["tables"]:
                    for table_data in pdf_data["tables"]:
                        page_num = table_data["page"]
                        table_content = table_data["content"]
                        
                        if table_content:
                            para = doc.add_paragraph()
                            run = para.add_run(f"\n【第{page_num}页表格】")
                            run.font.color.rgb = COLOR_WARNING
                            run.font.size = Pt(11)
                            run.bold = True
                            
                            for row in table_content:
                                # 对表格内容进行纠错
                                corrected_row, errors = self.corrector.correct_text(row)
                                if errors:
                                    run = para.add_run(f"\n【原文】{row}")
                                    run.font.color.rgb = COLOR_ERROR
                                    run = para.add_run(f"\n【建议】{corrected_row}")
                                    run.font.color.rgb = COLOR_SUCCESS
                                    all_errors.extend(errors)
                                else:
                                    run = para.add_run(f"\n{row}")
                                    run.font.color.rgb = COLOR_TEXT
                logger.info("✅ PDF表格处理完成")
            except Exception as e:
                logger.error(f"❌ PDF表格处理失败: {e}")
        
        # ===== 4. Word图片OCR =====
        if is_docx:
            logger.info("🖼 处理文档内图片...")
            try:
                image_paths = self.extract_images_from_docx(input_file)
                
                for img_path in image_paths:
                    text = extract_text_from_image(img_path)
                    
                    if text:
                        para = doc.add_paragraph()
                        run = para.add_run(f"\n【图片识别】\n{text}")
                        run.font.color.rgb = COLOR_INFO
                        
                        corrected, errors = self.corrector.correct_text(text)
                        
                        if errors:
                            run = para.add_run(f"\n【图片识别建议】{corrected}")
                            run.font.color.rgb = COLOR_SUCCESS
                            all_errors.extend(errors)
                logger.info("✅ 图片OCR处理完成")
            except Exception as e:
                logger.error(f"❌ 图片OCR处理失败: {e}")
        
        # PDF图片OCR处理
        if input_file.endswith(".pdf"):
            logger.info("🖼 处理PDF图片...")
            try:
                from pdf_processor import ocr_image
                if pdf_data and pdf_data["images"]:
                    for idx, img_data in enumerate(pdf_data["images"]):
                        page_num = img_data["page"]
                        image_bytes = img_data["image"]
                        
                        text = ocr_image(image_bytes)
                        
                        if text:
                            para = doc.add_paragraph()
                            run = para.add_run(f"\n【第{page_num}页图片{idx+1}识别】")
                            run.font.color.rgb = COLOR_WARNING
                            run.font.size = Pt(11)
                            run.bold = True
                            
                            run = para.add_run(f"\n{text}")
                            run.font.color.rgb = COLOR_INFO
                            
                            corrected, errors = self.corrector.correct_text(text)
                            
                            if errors:
                                run = para.add_run(f"\n【图片识别建议】{corrected}")
                                run.font.color.rgb = COLOR_SUCCESS
                                all_errors.extend(errors)
                logger.info("✅ PDF图片OCR处理完成")
            except Exception as e:
                logger.error(f"❌ PDF图片OCR处理失败: {e}")
        
        # ===== 5. 图片编号检测 =====
        logger.info("🔢 检测图片编号...")
        try:
            import re
            # 获取段落内容（兼容docx和PDF）
            doc_paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            
            image_errors = self.check_image_numbering(doc_paragraphs)
            all_errors.extend(image_errors)
            if image_errors:
                para = doc.add_paragraph()
                run = para.add_run(f"\n【图片编号问题】")
                run.font.color.rgb = COLOR_WARNING
                run.bold = True
                
                for e in image_errors:
                    run = para.add_run(f"\n → {e['message']}")
                    run.font.color.rgb = COLOR_ERROR
            logger.info(f"✅ 图片编号检测完成，发现 {len(image_errors)} 个问题")
        except Exception as e:
            logger.error(f"❌ 图片编号检测失败: {e}")
        
        # ===== 6. 表格编号检测 =====
        if is_docx:
            logger.info("📋 检测表格编号...")
            try:
                table_errors = self.check_table_numbering(doc)
                all_errors.extend(table_errors)
                if table_errors:
                    para = doc.add_paragraph()
                    run = para.add_run(f"\n【表格编号问题】")
                    run.font.color.rgb = COLOR_WARNING
                    run.bold = True
                    
                    for e in table_errors:
                        run = para.add_run(f"\n → {e['message']}")
                        run.font.color.rgb = COLOR_ERROR
                logger.info(f"✅ 表格编号检测完成，发现 {len(table_errors)} 个问题")
            except Exception as e:
                logger.error(f"❌ 表格编号检测失败: {e}")
        
        # ===== 7. 参考文献处理 =====
        logger.info("📚 检测参考文献...")
        try:
            # 获取段落内容（兼容docx和PDF）
            doc_paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            refs = extract_references(doc_paragraphs)
            logger.info(f"发现 {len(refs)} 条参考文献")
            
            # 检查编号连续性
            index_errors = check_index_sequence(refs)
            for err in index_errors:
                all_errors.append({
                    "type": "reference",
                    "level": "error",
                    "pos": None,
                    "text": "参考文献编号",
                    "suggestion": None,
                    "message": err
                })
            
            # 检查每条参考文献格式和内容
            for ref in refs:
                # 对参考文献内容进行文本纠错
                ref_corrected, ref_errors = self.corrector.correct_text(ref)
                if ref_errors:
                    all_errors.extend(ref_errors)
                    para = doc.add_paragraph()
                    run = para.add_run(f"\n【参考文献错别字】{ref}")
                    run.font.color.rgb = COLOR_ERROR
                    run = para.add_run(f"\n【修正后】{ref_corrected}")
                    run.font.color.rgb = COLOR_SUCCESS
                
                # 使用纠错后的文本进行格式检查
                check_ref = ref_corrected if ref_errors else ref
                
                # 检查参考文献格式
                errors = check_reference(check_ref)
                all_errors.extend(errors)
                
                if errors:
                    para = doc.add_paragraph()
                    run = para.add_run(f"\n【参考文献格式错误】{check_ref}")
                    run.font.color.rgb = COLOR_ERROR
                    
                    for e in errors:
                        run = para.add_run(f"\n → {e['message']}")
                        run.font.color.rgb = COLOR_WARNING
                    
                    run = para.add_run(f"\n【建议修改】{suggest_reference_fix(check_ref)}")
                    run.font.color.rgb = COLOR_SUCCESS
                    
            logger.info(f"✅ 参考文献检测完成，发现 {len([e for e in all_errors if e.get('type') == 'reference'])} 个问题")
            
        except Exception as e:
            logger.error(f"❌ 参考文献处理失败: {e}")
        
        # ===== 8. 在文档末尾添加统一的错误报告 =====
        logger.info("📝 生成统一错误报告...")
        try:
            if all_errors:
                # 添加分隔线
                para = doc.add_paragraph()
                run = para.add_run("\n" + "="*80)
                run.font.color.rgb = COLOR_INFO
                
                # 添加标题
                para = doc.add_paragraph()
                run = para.add_run("\n【智能检测报告 - 错误汇总】")
                run.font.color.rgb = RGBColor(0, 0, 128)  # 深蓝色
                run.font.size = Pt(14)
                run.bold = True
                
                # 按类型分组错误
                grouped_errors = {}
                for error in all_errors:
                    error_type = error.get('type', '其他')
                    if error_type not in grouped_errors:
                        grouped_errors[error_type] = []
                    grouped_errors[error_type].append(error)
                
                # 显示各类错误
                type_names = {
                    'spell': '错别字',
                    'grammar': '语法问题',
                    'semantic': '语义问题',
                    'reference': '参考文献格式',
                    'image': '图片编号',
                    'table': '表格编号'
                }
                
                for error_type, errors in grouped_errors.items():
                    type_name = type_names.get(error_type, error_type)
                    para = doc.add_paragraph()
                    run = para.add_run(f"\n【{type_name}】共 {len(errors)} 个问题")
                    run.font.color.rgb = COLOR_WARNING
                    run.bold = True
                    
                    for i, error in enumerate(errors, 1):
                        message = error.get('message', '')
                        text = error.get('text', '')
                        suggestion = error.get('suggestion', '')
                        
                        para = doc.add_paragraph()
                        run = para.add_run(f"\n  {i}. ")
                        run.font.color.rgb = COLOR_TEXT
                        
                        if text:
                            run = para.add_run(f"原文: {text}")
                            run.font.color.rgb = COLOR_ERROR
                        
                        if message:
                            run = para.add_run(f" → {message}")
                            run.font.color.rgb = COLOR_ERROR
                        
                        if suggestion:
                            run = para.add_run(f"\n    建议: {suggestion}")
                            run.font.color.rgb = COLOR_SUCCESS
                
                # 添加分隔线
                para = doc.add_paragraph()
                run = para.add_run("\n" + "="*80)
                run.font.color.rgb = COLOR_INFO
            
            logger.info("✅ 统一错误报告生成完成")
        except Exception as e:
            logger.error(f"❌ 生成错误报告失败: {e}")
        
        # ===== 9. 保存文档 =====
        logger.info(f"💾 保存结果到: {output_file}")
        try:
            os.makedirs(os.path.dirname(output_file), exist_ok=True)
            doc.save(output_file)
            logger.info("✅ 文件保存成功")
        except Exception as e:
            logger.error(f"❌ 保存文件失败: {e}")
            return None, []
        
        return output_file, all_errors
    
    def _build_correction_paragraph(self, para, text, corrected, errors, keep_images=False):
        """构建纠错段落（统一方法）
        
        Args:
            para: 段落对象
            text: 原始文本
            corrected: 修正后的文本
            errors: 错误列表
            keep_images: 是否保留段落中的图片（Word文档处理时使用）
        """
        if keep_images:
            # 保存段落中的图片（保留InlineShape对象）
            images = []
            for run in para.runs:
                # 检查run中是否包含图片
                for shape in run.element.xpath('.//w:drawing'):
                    images.append((run, shape))
            
            # 清空原段落
            for run in para.runs:
                run.text = ''
            
            # 如果段落中有图片，先添加图片
            for run, shape in images:
                # 创建新run并添加图片（保持原有图片）
                new_run = para.add_run()
                # 复制图片drawing元素到新run
                new_run.element.append(shape)
        
        # 添加原文（红色）
        run = para.add_run(f"【原文】")
        run.font.color.rgb = COLOR_ERROR
        run.font.size = Pt(11)
        run.bold = True
        
        # 高亮显示错误部分
        self._add_highlighted_text(para, text, errors, COLOR_ERROR)
        
        # 添加建议（绿色）
        run = para.add_run(f"\n【建议】")
        run.font.color.rgb = COLOR_SUCCESS
        run.font.size = Pt(11)
        run.bold = True
        
        # 添加修正后的文本
        run = para.add_run(corrected)
        run.font.color.rgb = COLOR_SUCCESS
        run.font.size = Pt(11)
        
        # 添加详细错误列表（蓝色）
        for e in errors:
            pos = e.get('pos', '未知')
            run = para.add_run(f"\n → 位置{pos}：{e['text']} → {e['suggestion']}")
            run.font.color.rgb = COLOR_INFO
            run.font.size = Pt(10)
    
    def _add_highlighted_text(self, para, text, errors, color):
        """高亮显示文本中的错误部分"""
        # 按位置排序错误
        sorted_errors = sorted(errors, key=lambda x: x.get('pos', 0))
        
        last_pos = 0
        for e in sorted_errors:
            pos = e.get('pos', 0)
            
            # 添加错误前的正常文本
            if pos > last_pos:
                run = para.add_run(text[last_pos:pos])
                run.font.color.rgb = COLOR_TEXT
                run.font.size = Pt(11)
            
            # 添加错误文本（高亮）
            run = para.add_run(e.get('text', ''))
            run.font.color.rgb = color
            run.font.size = Pt(11)
            run.bold = True
            
            last_pos = pos + len(e.get('text', ''))
        
        # 添加剩余的正常文本
        if last_pos < len(text):
            run = para.add_run(text[last_pos:])
            run.font.color.rgb = COLOR_TEXT
            run.font.size = Pt(11)
    
    def generate_report(self, errors):
        """生成检测报告"""
        report = generate_report(errors)
        details = generate_detail_report(errors)
        
        # 打印报告
        logger.info("\n" + "="*60)
        logger.info("📊 智能检测报告")
        logger.info("="*60)
        
        for k, v in report.items():
            logger.info(f"{k}: {v}")
        
        if details:
            logger.info("\n📄 错误详情：")
            for d in details:
                logger.info(f"  {d}")
        
        logger.info("="*60)
        
        return report, details

def main():
    """主函数"""
    parser = argparse.ArgumentParser(
        description="智能文本检测与格式批注系统 - 学术文档校验工具\n\n"
                    "支持的纠错模式:\n"
                    "  single: 单模型模式（仅macbert_finetuned，速度最快）\n"
                    "  dual: 双模型模式（macbert_finetuned + pycorrector，深度学习+规则引擎，推荐）"
    )
    parser.add_argument(
        "-i", "--input", 
        help="输入文件路径 (docx/pdf/png/jpg)",
        default=os.path.join(BASE_DIR, "test_files", "picture_error.jpg")
    )
    parser.add_argument(
        "-o", "--output",
        help="输出文件路径",
        default=os.path.join(BASE_DIR, "output", "picture_result.docx")
    )
    parser.add_argument(
        "-mode", "--mode",
        help="纠错模式: single（仅macbert）、dual（macbert+pycorrector双模型，推荐）",
        default=DEFAULT_MODE,
        choices=["single", "dual"]
    )
    
    args = parser.parse_args()
    
    # 验证输入文件
    if not os.path.exists(args.input):
        logger.error(f"❌ 输入文件不存在: {args.input}")
        sys.exit(1)
    
    # 创建校验器（支持单模型/双模型模式）
    checker = AcademicDocumentChecker(args.mode)
    
    # 加载模型
    if not checker.load_model():
        logger.error("❌ 模型加载失败，退出程序")
        sys.exit(1)
    
    # 处理文档
    output_path, errors = checker.process_document(args.input, args.output)
    
    if output_path:
        # 生成报告
        checker.generate_report(errors)
        logger.info(f"\n🎉 检测完成！结果已保存到: {output_path}")
    else:
        logger.error("❌ 处理失败")
        sys.exit(1)

if __name__ == "__main__":
    main()
